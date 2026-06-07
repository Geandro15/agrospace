import os

import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from PIL import Image
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import io, json, os, time
from pathlib import Path
from huggingface_hub import InferenceClient
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# Imports para análise de documentos e planilhas
try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# Imports para Web Scraping, XPath e Selenium
try:
    from lxml import etree
    LXML_OK = True
except ImportError:
    LXML_OK = False

try:
    from bs4 import BeautifulSoup
    BS4_OK = True
except ImportError:
    BS4_OK = False

import requests as req_lib

# ── Configuração da página ─────────────────────────────────────────────────
st.set_page_config(
    page_title="AgroSpace Dashboard",
    page_icon="🛰️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── CSS customizado ────────────────────────────────────────────────────────
st.markdown("""
<style>
    [data-testid="stSidebar"] { background: linear-gradient(180deg,#0d2137 0%,#0a3d1f 100%); }
    [data-testid="stSidebar"] * { color: #e8f5e9 !important; }
    .metric-card { background:#1a2744; border-radius:12px; padding:18px;
                   border-left:4px solid #40c463; margin:8px 0; }
    .metric-card h3 { color:#40c463; font-size:1.8em; margin:0; }
    .metric-card p  { color:#aaa; margin:0; font-size:0.85em; }
    .section-title  { color:#40c463; border-bottom:2px solid #40c463;
                      padding-bottom:6px; margin-bottom:16px; }
    .chat-user { background:#1e3a5f; border-radius:12px; padding:12px 16px;
                 margin:8px 0; border-left:3px solid #1f6feb; }
    .chat-bot  { background:#1a2f1a; border-radius:12px; padding:12px 16px;
                 margin:8px 0; border-left:3px solid #40c463; }
    .stButton>button { background:#40c463; color:#000; border:none;
                       border-radius:8px; font-weight:bold; }
    .stButton>button:hover { background:#2ea84d; color:#fff; }
    div[data-testid="metric-container"] { background:#1a2744; border-radius:10px; padding:10px; }
</style>
""", unsafe_allow_html=True)

# ── Inicializa RAG (cache para não recarregar) ─────────────────────────────
@st.cache_resource
def init_rag():
    DOCUMENTS = [
        {"titulo": "Sensoriamento Remoto na Agricultura de Precisão", "fonte": "Módulo 1",
         "texto": "O sensoriamento remoto por satélite revolucionou a agricultura de precisão. Satélites como Sentinel-2 (ESA) e Landsat-9 (NASA) capturam imagens multiespectrais que permitem calcular índices vegetativos como o NDVI (Normalized Difference Vegetation Index). O NDVI mede a saúde da vegetação comparando reflectância no infravermelho próximo (NIR) com o vermelho visível: NDVI = (NIR - RED) / (NIR + RED). Valores próximos de 1,0 indicam vegetação densa e saudável; valores abaixo de 0,2 indicam solo exposto ou vegetação estressada."},
        {"titulo": "IoT e Conectividade Satelital no Campo", "fonte": "Módulo 2",
         "texto": "A Internet das Coisas (IoT) agrícola integra sensores de solo, clima e equipamentos conectados a redes satelitais de baixa órbita (LEO). Constelações como Starlink (SpaceX), OneWeb e Telesat Lightspeed oferecem latência abaixo de 50ms e cobertura global. O protocolo LoRaWAN é amplamente adotado para comunicação de longa distância com baixo consumo energético em sensores distribuídos no campo."},
        {"titulo": "Inteligência Artificial e Machine Learning na Agricultura", "fonte": "Módulo 3",
         "texto": "Modelos de Machine Learning aplicados à agricultura espacial incluem: CNN para classificação de culturas, LSTM para previsão de séries temporais, Random Forest e XGBoost para estimativa de produtividade, YOLO e Faster R-CNN para identificação de pragas em imagens de drone. O Google Earth Engine permite processar petabytes de imagens satelitais na nuvem."},
        {"titulo": "Drones e Veículos Aéreos Não Tripulados (VANT)", "fonte": "Módulo 4",
         "texto": "Drones agrícolas operam em baixa altitude (30–120m) com resolução centimétrica. Modelos populares: DJI Agras T40 (pulverização) e senseFly eBee (mapeamento). Geram ortomosaicos, modelos digitais de elevação, mapas NDVI e nuvens de pontos 3D. Reduzem uso de agroquímicos em até 30-40%."},
        {"titulo": "Economia Espacial e o Agronegócio Brasileiro", "fonte": "Módulo 5",
         "texto": "A nova economia espacial movimentou US$ 546 bilhões em 2023, com projeções de US$ 1 trilhão até 2030. O Brasil possui o Centro de Lançamento de Alcântara (CLA) e opera o satélite CBERS. Startups agtech captaram mais de R$ 2,5 bilhões entre 2020 e 2024."},
        {"titulo": "Irrigação Inteligente Baseada em Dados Satelitais", "fonte": "Módulo 6",
         "texto": "A irrigação responde por 70% do consumo global de água doce. A irrigação inteligente reduz o desperdício em até 50%. O índice ET calculado por dados de satélite permite estimar a demanda hídrica real. Sistemas IoT ajustam automaticamente a lâmina d'água com base em umidade do solo, previsão meteorológica e NDWI via satélite."},
        {"titulo": "Monitoramento de Desmatamento e Rastreabilidade", "fonte": "Módulo 7",
         "texto": "O Brasil monitora o desmatamento com PRODES, DETER e MapBiomas. A UE exige certificação de não desmatamento (EUDR) para soja, carne bovina e cacau. O CAR integrado a dados satelitais permite fiscalização automatizada do Código Florestal."},
        {"titulo": "Modelos de Previsão de Safra", "fonte": "Módulo 8",
         "texto": "Modelos de previsão integram dados satelitais, climáticos e agronômicos. O DSSAT simula crescimento de culturas. A CONAB usa MODIS/NDVI para estimativas mensais. IA supervisionada atinge acurácia acima de 90% na previsão de soja e milho com 2-3 meses de antecedência."},
        {"titulo": "Blockchain e Contratos Inteligentes no Agro", "fonte": "Módulo 9",
         "texto": "Blockchain com dados satelitais cria rastreabilidade imutável na cadeia agroalimentar. Smart contracts automatizam pagamentos por créditos de carbono com verificação via satélite. O crédito de carbono voluntário foi negociado a US$ 8-15 por tonelada de CO2 em 2023."},
        {"titulo": "Geotecnologias e SIG", "fonte": "Módulo 10",
         "texto": "SIG integram camadas espaciais para tomada de decisão agrícola. QGIS e ArcGIS permitem cruzamento de dados de solo, clima e sensoriamento remoto. O ZARC define janelas de plantio de baixo risco para 40+ culturas. GNSS de alta precisão (RTK, PPP) é base para autopiloto de maquinários agrícolas."}
    ]
    docs = [Document(page_content=d["texto"], metadata={"titulo": d["titulo"], "fonte": d["fonte"]}) for d in DOCUMENTS]
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=80)
    chunks = splitter.split_documents(docs)
    emb = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        model_kwargs={"device": "cpu"}, encode_kwargs={"normalize_embeddings": True}
    )
    vs = FAISS.from_documents(chunks, emb)
    llm = InferenceClient(provider="novita", api_key=os.environ.get("HF_TOKEN", ""))
    return vs, llm, DOCUMENTS

def gerar_resposta_rag(pergunta, vectorstore, client):
    docs = vectorstore.similarity_search(pergunta, k=3)
    contexto = "\\n\\n".join([f"[{d.metadata['titulo']}]\\n{d.page_content}" for d in docs])
    fontes = list({d.metadata["titulo"] for d in docs})
    resp = client.chat.completions.create(
        model="meta-llama/Llama-3.1-8B-Instruct",
        messages=[
            {"role": "system", "content": "Você é um assistente especialista em agricultura inteligente e nova economia espacial. Responda em português, de forma clara, com base APENAS no contexto fornecido."},
            {"role": "user", "content": f"CONTEXTO:\\n{contexto}\\n\\nPERGUNTA: {pergunta}"}
        ],
        max_tokens=512, temperature=0.3
    )
    return resp.choices[0].message.content.strip(), fontes

# ── Sidebar ────────────────────────────────────────────────────────────────
st.sidebar.markdown("""
<div style='text-align:center; padding:10px 0 6px 0;'>
  <div style='font-size:2.2em;'>🛰️</div>
  <div style='font-size:1.3em; font-weight:bold; color:#40c463;'>AgroSpace</div>
  <div style='font-size:0.8em; color:#a5d6a7;'>Nova Economia Espacial</div>
  <div style='font-size:0.75em; color:#81c784;'>Agricultura Inteligente</div>
</div>
""", unsafe_allow_html=True)
st.sidebar.markdown("---")
st.sidebar.markdown("---")

menu = st.sidebar.radio("📋 Menu", [
    "🏠 Home",
    "📡 Sensoriamento Remoto",
    "🤖 IA & Machine Learning",
    "🌊 Irrigação Inteligente",
    "🚁 Drones & VANTs",
    "🌳 Sustentabilidade & Carbono",
    "🇧🇷 Economia Espacial Brasil",
    "🖼️ Análise de Imagem NDVI",
    "🔄 Automação Robótica de Processos",
    "📄 Pesquisa Doc",
    "📊 Análise de Planilhas IA",
    "🌐 Módulo de coleta automatizada de dados",
    "💬 Chat AgroSpace"
])

st.sidebar.markdown("---")
st.sidebar.markdown("**📚 Fontes de Pesquisa:**")
st.sidebar.markdown("""
<small>
🔬 <a href='https://earthdata.nasa.gov' target='_blank'>NASA — MODIS, Landsat</a><br>
🛰️ <a href='https://www.copernicus.eu' target='_blank'>ESA Copernicus — Sentinel-2</a><br>
🌱 <a href='https://www.embrapa.br' target='_blank'>Embrapa — IRRIGAR, ZARC</a><br>
🌳 <a href='https://www.inpe.br' target='_blank'>INPE — PRODES, DETER</a><br>
🌍 <a href='https://cropmonitor.org' target='_blank'>FAO — CROP MONITOR</a><br>
🌊 <a href='https://www.noaa.gov' target='_blank'>NOAA — El Niño/ENSO</a><br>
💻 <a href='https://github.com/satellite-image-deep-learning' target='_blank'>GitHub — Satellite ML</a><br>
📊 <a href='https://www.kaggle.com/competitions/csiro-biomass' target='_blank'>Kaggle — CSIRO Biomass</a><br>
🤗 <a href='https://huggingface.co' target='_blank'>HuggingFace — Llama/MiniLM</a><br>
🌐 <a href='https://earthengine.google.com' target='_blank'>Google Earth Engine</a><br>
🗺️ <a href='https://terrabrasilis.dpi.inpe.br' target='_blank'>TerraBrasilis — INPE</a><br>
🧑‍🔬 <a href='https://colab.research.google.com/drive/1jciJ0UGDUnWwZDPB7aRTKZjlq6EbQQKb?usp=sharing' target='_blank'>RAG Agricultura Espacial — Geandro Dezordi</a>
</small>
""", unsafe_allow_html=True)
st.sidebar.markdown("---")
st.sidebar.markdown("**🛰️ Satélites para Agricultura:**")
st.sidebar.markdown("""
<small>
🌍 <a href='https://earth.google.com/web/' target='_blank'>Google Earth Web</a><br>
🗺️ <a href='https://maps.google.com/' target='_blank'>Google Maps — Satélite</a><br>
🔭 <a href='https://worldview.earthdata.nasa.gov/' target='_blank'>NASA Worldview</a><br>
🌀 <a href='https://zoom.earth/' target='_blank'>Zoom Earth</a><br>
📡 <a href='https://apps.sentinel-hub.com/eo-browser/' target='_blank'>Sentinel Hub EO Browser</a><br>
🇪🇺 <a href='https://browser.dataspace.copernicus.eu/' target='_blank'>Copernicus Browser</a>
</small>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# 🏠 HOME
# ══════════════════════════════════════════════════════════════════
if menu == "🏠 Home":
    st.title("🌾🛰️ AgroSpace Dashboard")
    st.markdown("### Nova Economia Espacial & Agricultura Inteligente")
    st.markdown("---")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("💰 Mercado Espacial 2023", "US$ 546 bi", "+15% a.a.")
    with col2:
        st.metric("🌾 Projeção 2030", "US$ 1 trilhão", "+83%")
    with col3:
        st.metric("🇧🇷 Agtech Brasil", "R$ 2,5 bi", "2020–2024")
    with col4:
        st.metric("💧 Economia de Água", "até 50%", "Irrigação inteligente")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 📈 Crescimento do Mercado Espacial")
        anos = [2015,2016,2017,2018,2019,2020,2021,2022,2023,2024,2025,2026,2027,2028,2029,2030]
        valores = [330,350,385,415,440,447,470,499,546,600,660,730,800,875,935,1000]
        df_mkt = pd.DataFrame({"Ano": anos, "US$ Bilhões": valores})
        fig = px.area(df_mkt, x="Ano", y="US$ Bilhões",
                      color_discrete_sequence=["#40c463"],
                      template="plotly_dark")
        fig.add_vline(x=2026, line_dash="dash", line_color="#e3b341",
                      annotation_text="Hoje", annotation_position="top right")
        fig.update_layout(margin=dict(l=0,r=0,t=10,b=0), height=280)
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 🥧 Distribuição por Setor")
        setores = ["Satélites Comerciais","Lançamento","Serviços de Solo","Observação da Terra","Agtech Espacial"]
        valores_s = [38, 22, 18, 14, 8]
        fig2 = px.pie(values=valores_s, names=setores, hole=0.4,
                      color_discrete_sequence=px.colors.sequential.Greens_r,
                      template="plotly_dark")
        fig2.update_layout(margin=dict(l=0,r=0,t=10,b=0), height=280)
        st.plotly_chart(fig2, use_container_width=True)

    st.markdown("---")
    st.markdown("#### 🗺️ Módulos deste Dashboard")
    modulos = [
        ("📡","Sensoriamento Remoto","NDVI, Sentinel-2, Landsat, índices vegetativos"),
        ("🤖","IA & Machine Learning","CNN, LSTM, Random Forest, previsão de safra"),
        ("🌊","Irrigação Inteligente","ET, NDWI, IoT, redução de 50% no consumo hídrico"),
        ("🚁","Drones & VANTs","DJI Agras, senseFly eBee, fotogrametria, Pix4D"),
        ("🌳","Sustentabilidade","Créditos de carbono, PRODES, EUDR, MapBiomas"),
        ("🇧🇷","Economia Brasil","CLA Alcântara, CBERS, R$ 2,5 bi em agtech"),
        ("🖼️","Análise NDVI","Upload de imagem e diagnóstico agronômico automático"),
        ("🔄","Pipeline RPA","Fluxo de automação inteligente inspirado na exploração espacial"),
        ("📄","Documentos IA","Upload de PDF/Word com análise automática por IA"),
        ("📊","Planilhas IA","Upload de CSV/XLSX com gráficos e insights automáticos"),
        ("🌐","Web Scraping & XPath","Scraping de portais, parsing XML/KML e automação Selenium"),
        ("💬","Chat AgroSpace","Assistente IA com base nos documentos do projeto"),
    ]
    cols = st.columns(4)
    for i, (icon, titulo, desc) in enumerate(modulos):
        with cols[i % 4]:
            st.markdown(f"""
            <div style='background:#1a2744;border-radius:10px;padding:14px;margin:6px 0;border-left:3px solid #40c463'>
            <b style='color:#40c463'>{icon} {titulo}</b><br>
            <small style='color:#aaa'>{desc}</small></div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════
# 📡 SENSORIAMENTO REMOTO
# ══════════════════════════════════════════════════════════════════
elif menu == "📡 Sensoriamento Remoto":
    st.title("📡 Sensoriamento Remoto na Agricultura")
    st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1: st.metric("🛰️ Satélites Ativos", "800+", "2024")
    with col2: st.metric("📏 Resolução Sentinel-2", "10 metros", "por pixel")
    with col3: st.metric("🔄 Revisita", "5 dias", "Sentinel-2")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 📊 Índices Vegetativos — Escala de Valores")
        indices = ["NDVI","EVI","NDWI","NDRE","SAVI","GNDVI"]
        v_min   = [-1.0,-1.0,-1.0,-1.0,-1.0,-1.0]
        v_max   = [1.0, 1.0, 1.0, 1.0, 1.0, 1.0]
        v_ideal = [0.7, 0.6, 0.3, 0.5, 0.6, 0.65]
        df_idx = pd.DataFrame({"Índice":indices,"Valor Ideal":v_ideal})
        fig = px.bar(df_idx, x="Valor Ideal", y="Índice", orientation="h",
                     color="Valor Ideal", color_continuous_scale="Greens",
                     template="plotly_dark", title="Valor Ideal por Índice")
        fig.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=300)
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        cultura_sel = st.radio("Selecione a cultura:", ["🌱 Soja", "🎋 Cana-de-açúcar"], horizontal=True)
        semanas = list(range(1, 21))
        np.random.seed(10)
        if cultura_sel == "🌱 Soja":
            ndvi_serie = np.concatenate([
                np.linspace(0.1, 0.85, 10) + np.random.normal(0, 0.02, 10),
                np.linspace(0.85, 0.2, 10) + np.random.normal(0, 0.02, 10)
            ])
            titulo_serie = "Evolução NDVI — Soja (20 semanas)"
            cor_serie    = "#40c463"
            desc_serie   = "Soja: crescimento rápido nas primeiras 10 semanas (NDVI até 0,85), seguido de senescência na colheita."
        else:
            ndvi_serie = np.concatenate([
                np.linspace(0.05, 0.45, 5) + np.random.normal(0, 0.015, 5),
                np.linspace(0.45, 0.80, 8) + np.random.normal(0, 0.02, 8),
                np.linspace(0.80, 0.55, 7) + np.random.normal(0, 0.02, 7)
            ])
            titulo_serie = "Evolução NDVI — Cana-de-açúcar (20 semanas)"
            cor_serie    = "#f4c430"
            desc_serie   = "Cana: crescimento lento inicial, pico de biomassa entre semanas 13-15 (NDVI ~0,80) e leve queda pré-colheita."

        df_ndvi = pd.DataFrame({"Semana": semanas, "NDVI": ndvi_serie})
        fig2 = px.line(df_ndvi, x="Semana", y="NDVI", markers=True,
                       color_discrete_sequence=[cor_serie],
                       template="plotly_dark", title=titulo_serie)
        fig2.add_hline(y=0.5, line_dash="dash", line_color="yellow",
                       annotation_text="Limiar saudável")
        fig2.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=270)
        st.plotly_chart(fig2, use_container_width=True)
        st.caption(desc_serie)

    # Expander de explicação dos índices — abre ao clicar
    with st.expander("📖 Clique aqui para entender cada índice vegetativo do gráfico"):
        st.markdown("""
| Índice | Nome completo | O que mede | Valor ideal | Uso principal |
|--------|--------------|------------|-------------|---------------|
| **NDVI** | Normalized Difference Vegetation Index | Saúde geral da vegetação | ≥ 0,7 | Monitoramento de culturas, detecção de estresse |
| **EVI** | Enhanced Vegetation Index | Vigor vegetativo em dosséis densos | ≥ 0,6 | Cerrado, Amazônia — menos saturação que NDVI |
| **NDWI** | Normalized Difference Water Index | Conteúdo de água na vegetação | ≥ 0,3 | Estresse hídrico, irrigação de precisão |
| **NDRE** | Normalized Difference Red Edge | Teor de clorofila e nitrogênio | ≥ 0,5 | Fertilidade, deficiência nutricional precoce |
| **SAVI** | Soil Adjusted Vegetation Index | Vegetação em áreas com solo exposto | ≥ 0,6 | Regiões áridas, início do ciclo da cultura |
| **GNDVI** | Green Normalized Difference Vegetation Index | Atividade fotossintética (clorofila) | ≥ 0,65 | Maturidade da cultura, estimativa de produtividade |

**Como interpretar:** valores próximos de 1,0 indicam vegetação densa e saudável;
valores abaixo de 0,2 indicam solo exposto ou vegetação morta.
Fórmula base do NDVI: `(NIR - RED) / (NIR + RED)` — proposta por Rouse et al. (1974).
        """)

    st.markdown("#### 🛰️ Satélites de Observação Agrícola")
    df_sat = pd.DataFrame({
        "Satélite": ["Sentinel-2A/B","Landsat-9","CBERS-4A","Planet Dove","MODIS (Terra/Aqua)"],
        "Agência":  ["ESA","NASA/USGS","INPE/CAST","Planet Labs","NASA"],
        "Resolução":["10m","30m","8m","3m","250m"],
        "Revisita": ["5 dias","16 dias","5 dias","Diário","1–2 dias"],
        "Uso Agro": ["NDVI, EVI, NDWI","Histórico longo","Monitoramento BR","Alta frequência","Escala global"]
    })
    st.dataframe(df_sat, use_container_width=True, hide_index=True)

# ══════════════════════════════════════════════════════════════════
# 🤖 IA & MACHINE LEARNING
# ══════════════════════════════════════════════════════════════════
elif menu == "🤖 IA & Machine Learning":
    st.title("🤖 IA & Machine Learning na Agricultura")
    st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1: st.metric("🎯 Acurácia Previsão Safra", "90%+", "Redes neurais")
    with col2: st.metric("⏱️ Antecedência", "2–3 meses", "Antes da colheita")
    with col3: st.metric("📉 Redução Defensivos", "30–40%", "Com drones + IA")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 🧠 Modelos de ML por Aplicação")
        modelos = ["CNN","LSTM","Random Forest","XGBoost","YOLO","U-Net","GPR"]
        acuracia = [92, 88, 91, 93, 89, 94, 87]
        aplicacao= ["Classif. Culturas","Previsão Temporal","Produtividade",
                    "Produtividade","Detecção Pragas","Segmentação","Previsão Safra"]
        df_ml = pd.DataFrame({"Modelo":modelos,"Acurácia (%)":acuracia,"Aplicação":aplicacao})
        fig = px.bar(df_ml, x="Modelo", y="Acurácia (%)", color="Acurácia (%)",
                     hover_data=["Aplicação"], color_continuous_scale="Greens",
                     template="plotly_dark", title="Acurácia por Modelo de IA")
        fig.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=300)
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 📊 Previsão de Produtividade — Soja (ton/ha)")
        anos = list(range(2018, 2031))
        real = [3.1,3.3,3.0,3.5,3.4,3.7,3.8,None,None,None,None,None,None]
        previsto = [None,None,None,None,None,3.7,3.9,4.1,4.3,4.4,4.6,4.8,5.0]
        df_prod = pd.DataFrame({"Ano":anos,"Real":real,"Previsto (IA)":previsto})
        fig2 = go.Figure()
        fig2.add_trace(go.Scatter(x=anos, y=real, name="Real", line=dict(color="#40c463",width=2)))
        fig2.add_trace(go.Scatter(x=anos, y=previsto, name="Previsto (IA)",
                                  line=dict(color="#1f6feb",width=2,dash="dot")))
        fig2.update_layout(template="plotly_dark", height=300,
                           margin=dict(l=0,r=0,t=10,b=0),
                           title="Produtividade Soja — Real vs Previsto por IA")
        st.plotly_chart(fig2, use_container_width=True)

    st.markdown("#### 🔬 Pipeline de IA Agrícola — clique em uma etapa para saber mais")

    PIPELINE_INFO = {
        "🛰️ Imagem Satélite/Drone": {
            "desc": "Coleta de dados brutos via sensores multiespectrais.",
            "numeros": "Sentinel-2 gera ~1 TB/dia. Planet Labs captura 3 milhões de km²/dia.",
            "conclusao": "Investir em acesso a dados satelitais é o ponto de partida. O custo de 1 cena Sentinel-2 caiu de US$ 20 (2010) para zero (2024) — tornando a análise acessível a qualquer produtor."
        },
        "⚙️ Pré-processamento": {
            "desc": "Correção atmosférica, remoção de nuvens, normalização de bandas.",
            "numeros": "Reduz ruído em até 40%. Sem pré-processamento, acurácia dos modelos cai de 90% para 60%.",
            "conclusao": "Vale cada centavo: sem pré-processamento correto, qualquer modelo de IA produz resultados incorretos. É a etapa mais crítica e muitas vezes mais negligenciada do pipeline."
        },
        "📐 Extração de Features": {
            "desc": "Cálculo de NDVI, EVI, NDWI e outras variáveis derivadas das imagens.",
            "numeros": "Um único pixel Sentinel-2 gera até 13 bandas e 20+ índices derivados.",
            "conclusao": "Extrair os índices certos multiplica a capacidade preditiva. NDRE supera NDVI em 15% na detecção de deficiência de nitrogênio — um dado que pode economizar R$ 200/ha em fertilizantes."
        },
        "🤖 Modelo de ML/DL": {
            "desc": "Treinamento de CNN, LSTM, XGBoost ou Random Forest com os dados extraídos.",
            "numeros": "XGBoost atinge 93% de acurácia. CNN + LSTM combinados chegam a 96% em previsão de safra.",
            "conclusao": "Modelos de IA pagam seu custo de desenvolvimento em 1 safra: uma previsão 90% precisa com 3 meses de antecedência permite negociar contratos futuros com vantagem de R$ 50-150/ton."
        },
        "📊 Predição": {
            "desc": "Geração de mapas de produtividade, alertas de pragas e recomendações de manejo.",
            "numeros": "Modelos bem treinados reduzem perdas por pragas em 25% e por clima em 18%.",
            "conclusao": "A predição transformou a agricultura reativa em proativa. Detectar uma praga 2 semanas antes do dano visível reduz o custo de controle em até 60% e preserva a produtividade."
        },
        "✅ Decisão Agronômica": {
            "desc": "Ação baseada nas predições: aplicação localizada, ajuste de irrigação, manejo diferenciado.",
            "numeros": "Agricultores que usam decisões baseadas em IA têm ROI médio de 320% sobre o investimento em tecnologia.",
            "conclusao": "Este é o elo que transforma dados em lucro. A decisão agronômica assistida por IA reduz custos operacionais em 15-30% e aumenta a produtividade em 10-25% — justificando completamente o investimento em todo o pipeline."
        },
    }

    etapas_labels = list(PIPELINE_INFO.keys())
    cols_pipe = st.columns(len(etapas_labels))
    for i, etapa in enumerate(etapas_labels):
        with cols_pipe[i]:
            if st.button(etapa, key=f"pipe_{i}", use_container_width=True):
                st.session_state["etapa_sel"] = etapa

    if "etapa_sel" in st.session_state:
        info = PIPELINE_INFO[st.session_state["etapa_sel"]]
        st.markdown(f"""
<div style='background:#1a2744; border-left:4px solid #40c463; border-radius:10px;
            padding:18px; margin:12px 0;'>
  <h4 style='color:#40c463; margin:0 0 8px 0;'>{st.session_state["etapa_sel"]}</h4>
  <p style='color:#ccc; margin:4px 0;'><b>O que é:</b> {info["desc"]}</p>
  <p style='color:#40c463; margin:4px 0;'><b>Números:</b> {info["numeros"]}</p>
  <p style='color:#e8f5e9; margin:8px 0 0 0;'><b>Por que vale investir:</b> {info["conclusao"]}</p>
</div>
        """, unsafe_allow_html=True)

    # Funil visual abaixo dos botões
    fig3 = go.Figure(go.Funnel(
        y=["Imagem Satélite/Drone","Pré-processamento","Extração de Features",
           "Modelo de ML/DL","Predição","Decisão Agronômica"],
        x=[100,95,85,80,78,75],
        marker_color=["#40c463","#2ea84d","#1f8a3a","#1f6feb","#1a5ec4","#e3b341"]
    ))
    fig3.update_layout(template="plotly_dark", height=260, margin=dict(l=0,r=0,t=10,b=0))
    st.plotly_chart(fig3, use_container_width=True)

# ══════════════════════════════════════════════════════════════════
# 🌊 IRRIGAÇÃO INTELIGENTE
# ══════════════════════════════════════════════════════════════════
elif menu == "🌊 Irrigação Inteligente":
    st.title("🌊 Irrigação Inteligente por Satélite")
    st.markdown("---")

    col1,col2,col3 = st.columns(3)
    with col1: st.metric("💧 Consumo Agrícola Global","70%","da água doce")
    with col2: st.metric("📉 Redução com IoT+Satélite","até 50%","no desperdício")
    with col3: st.metric("🌡️ Evapotranspiração","ET calculada","via SEBAL/METRIC")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 💧 Consumo Hídrico: Tradicional vs Inteligente")
        culturas = ["Soja","Milho","Cana-de-açúcar","Café","Algodão"]
        trad = [550,600,1800,900,700]
        intel = [290,330,950,500,380]
        df_agua = pd.DataFrame({"Cultura":culturas,"Tradicional (mm)":trad,"Inteligente (mm)":intel})
        fig = go.Figure()
        fig.add_bar(x=culturas, y=trad, name="Tradicional", marker_color="#e74c3c")
        fig.add_bar(x=culturas, y=intel, name="Inteligente", marker_color="#40c463")
        fig.update_layout(barmode="group", template="plotly_dark",
                          height=300, margin=dict(l=0,r=0,t=10,b=0),
                          title="Lâmina d'água aplicada (mm/ciclo)")
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 📡 Índices Hídricos por Satélite")
        meses = ["Jan","Fev","Mar","Abr","Mai","Jun","Jul","Ago","Set","Out","Nov","Dez"]
        ndwi = [0.3,0.25,0.4,0.35,0.2,0.1,0.05,0.08,0.15,0.28,0.35,0.32]
        et   = [5.5,5.2,4.8,3.9,3.1,2.5,2.2,2.8,3.5,4.5,5.0,5.4]
        fig2 = make_subplots(specs=[[{"secondary_y": True}]])
        fig2.add_trace(go.Bar(x=meses, y=et, name="ET (mm/dia)", marker_color="#1f6feb"), secondary_y=False)
        fig2.add_trace(go.Scatter(x=meses, y=ndwi, name="NDWI", line=dict(color="#40c463",width=2),mode="lines+markers"), secondary_y=True)
        fig2.update_layout(template="plotly_dark", height=300, margin=dict(l=0,r=0,t=10,b=0))
        st.plotly_chart(fig2, use_container_width=True)

# ══════════════════════════════════════════════════════════════════
# 🚁 DRONES & VANTs
# ══════════════════════════════════════════════════════════════════
elif menu == "🚁 Drones & VANTs":
    st.title("🚁 Drones & VANTs na Agricultura")
    st.markdown("---")

    col1,col2,col3 = st.columns(3)
    with col1: st.metric("📐 Resolução máx.","1–2 cm/pixel","câmeras multiespectrais")
    with col2: st.metric("📉 Redução Defensivos","30–40%","aplicação localizada")
    with col3: st.metric("💰 Mercado Drones Agro","US$ 11,2 bi","projeção 2027")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 🚁 Comparativo de Modelos")
        df_drones = pd.DataFrame({
            "Modelo":   ["DJI Agras T40","senseFly eBee X","DJI Phantom 4 Multi","Parrot Sequoia+","AgDrone"],
            "Tipo":     ["Pulverização","Mapeamento","Multiespectral","Sensor","Mapeamento"],
            "Autonomia":["10 min","90 min","30 min","N/A","60 min"],
            "Área/dia": ["40 ha","500 ha","100 ha","N/A","200 ha"],
            "Câmera":   ["RGB","RGB+Multi","5 bandas","4 bandas","RGB+Multi"]
        })
        st.dataframe(df_drones, use_container_width=True, hide_index=True)

    with col_b:
        st.markdown("#### 📈 Crescimento do Mercado de Drones Agro")
        anos_d = [2019,2020,2021,2022,2023,2024,2025,2026,2027]
        mkt_d  = [1.2, 1.5, 2.1, 3.0, 4.5, 6.0, 7.8, 9.5, 11.2]
        df_mkt_d = pd.DataFrame({"Ano":anos_d,"US$ Bilhões":mkt_d})
        fig = px.bar(df_mkt_d, x="Ano", y="US$ Bilhões",
                     color="US$ Bilhões", color_continuous_scale="Greens",
                     template="plotly_dark", title="Mercado Global de Drones Agrícolas")
        fig.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=280)
        st.plotly_chart(fig, use_container_width=True)

# ══════════════════════════════════════════════════════════════════
# 🌳 SUSTENTABILIDADE & CARBONO
# ══════════════════════════════════════════════════════════════════
elif menu == "🌳 Sustentabilidade & Carbono":
    st.title("🌳 Sustentabilidade & Mercado de Carbono")
    st.markdown("---")

    col1,col2,col3 = st.columns(3)
    with col1: st.metric("💰 Preço Carbono Agro","US$ 8–15","por tonelada CO₂")
    with col2: st.metric("🎯 Mercado 2030","US$ 50 bi","créditos voluntários")
    with col3: st.metric("🌿 EUDR","Em vigor","desde 2023")

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 🌱 Práticas de Sequestro de Carbono")
        praticas = ["Plantio Direto","ILPF","Recuperação Pastagens","Reflorestamento","Biochar"]
        tonco2   = [1.5, 3.2, 2.1, 5.0, 2.8]
        df_carb  = pd.DataFrame({"Prática":praticas,"t CO₂/ha/ano":tonco2})
        fig = px.bar(df_carb, x="t CO₂/ha/ano", y="Prática", orientation="h",
                     color="t CO₂/ha/ano", color_continuous_scale="Greens",
                     template="plotly_dark", title="Sequestro de Carbono por Prática")
        fig.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=300)
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 📊 Desmatamento Amazônia — PRODES/INPE (km²/ano)")
        anos_d  = [2020, 2021, 2022, 2023, 2024, 2025]
        desmata = [10851, 13038, 11594, 9064, 6518, 5796]
        df_dmt  = pd.DataFrame({"Ano": anos_d, "km² desmatados": desmata})
        fig2 = px.line(df_dmt, x="Ano", y="km² desmatados", markers=True,
                       color_discrete_sequence=["#e74c3c"],
                       template="plotly_dark", title="Desmatamento Amazônia Legal — PRODES/INPE")
        fig2.add_annotation(x=2025, y=5796,
                            text="5.796 km²<br>Menor em 11 anos",
                            showarrow=True, arrowhead=2,
                            arrowcolor="#40c463", font=dict(color="#40c463", size=11),
                            bgcolor="#0d2137", bordercolor="#40c463")
        fig2.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=300)
        st.plotly_chart(fig2, use_container_width=True)

        with st.expander("📋 Dados completos e fontes PRODES/INPE"):
            st.markdown("""
**Série histórica recente — Amazônia Legal (km²):**

| Ano | Desmatamento (km²) | Variação |
|-----|--------------------|---------|
| 2025 | **5.796** (estimativa) | -11,1% |
| 2024 | 6.518 | -28,2% |
| 2023 | 9.064 | -21,8% |
| 2022 | 11.594 | -11,1% |
| 2021 | 13.038 | +20,2% |
| 2020 | 10.851 | — |

2025 marca o **4° ano consecutivo de queda** e o menor valor dos últimos 11 anos.
O período PRODES vai de **1° de agosto** de um ano a **31 de julho** do seguinte.

**Fontes:**
- [Nota Técnica PRODES 2025 — INPE](https://data.inpe.br/biomasbr/wp-content/uploads/sites/3/2025/10/20251015Nota_tecnica_EstimativaPRODES_2025_F.pdf)
- [TerraBrasilis — Visualizador Interativo](https://terrabrasilis.dpi.inpe.br/app/dashboard/deforestation/biomes/legal_amazon/increments)
- [G1 — Amazônia: desmatamento 2025](https://g1.globo.com/meio-ambiente/noticia/2025/10/30/brasil-desmatamento.ghtml)
            """)

# ══════════════════════════════════════════════════════════════════
# 🇧🇷 ECONOMIA ESPACIAL BRASIL
# ══════════════════════════════════════════════════════════════════
elif menu == "🇧🇷 Economia Espacial Brasil":
    st.title("🇧🇷 Economia Espacial Brasileira")
    st.markdown("---")

    col1,col2,col3 = st.columns(3)
    with col1: st.metric("🚀 CLA Alcântara","2° melhor","economia 30% combustível")
    with col2: st.metric("🛰️ Satélites BR","CBERS + SGDC","operacionais")
    with col3: st.metric("🌱 Agtech captações","R$ 2,5 bi","2020–2024")

    st.markdown("---")

    with st.expander("🚀 Por que o CLA Alcântara é estratégico? — Clique para saber"):
        st.markdown("""
O **Centro de Lançamento de Alcântara (CLA)** é considerado um dos melhores espaçoportos
do mundo, superando o Cabo Canaveral (EUA) e o Cosmódromo de Baikonur (Cazaquistão)
em eficiência energética e economia de custos.

**Vantagens competitivas:**

- **Economia de combustível:** Fica a apenas 2° ao sul do Equador, onde a velocidade
  de rotação da Terra é máxima (~465 m/s). Isso impulsiona o foguete naturalmente,
  gerando economia de até **30% de propelente**.

- **Maior carga útil:** A economia de combustível permite transportar cargas mais pesadas
  para a mesma órbita — vantagem direta para satélites agrícolas e de observação.

- **Clima favorável:** Não está em rota de furacões (como nos EUA) e não sofre com
  frio extremo. O clima quente reduz a densidade do ar e o atrito aerodinâmico.

- **Segurança:** Localizado em uma península com baixa densidade populacional,
  voltado para o Oceano Atlântico — minimiza riscos e facilita recuperação de estágios.

**Fontes:**
[Wikipedia — CLA](https://pt.wikipedia.org/wiki/Centro_Espacial_de_Alc%C3%A2ntara) |
[TecMundo — Base de Alcântara](https://www.tecmundo.com.br/ciencia/261195-base-alcantara-melhores-mundo-lancar-foguetes.htm) |
[SpaceLab — Por que a melhor base é brasileira?](https://www.youtube.com/watch?v=OxBz9LAEgpQ)
        """)

    st.markdown("---")
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### 💰 Investimento em Agtech no Brasil")
        anos_ag = [2018,2019,2020,2021,2022,2023,2024]
        invest  = [0.3, 0.5, 0.8, 1.2, 1.8, 2.2, 2.5]
        df_ag = pd.DataFrame({"Ano":anos_ag,"R$ Bilhões":invest})
        fig = px.area(df_ag, x="Ano", y="R$ Bilhões",
                      color_discrete_sequence=["#009c3b"],
                      template="plotly_dark", title="Captações Agtech — Brasil")
        fig.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=280)
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 🌾 Produção Agro Brasileira — Principais Culturas")
        culturas_br = ["Soja","Milho","Cana-de-açúcar","Café","Algodão","Laranja"]
        producao_mt = [162, 137, 715, 3.5, 6.8, 16]
        fig2 = px.bar(x=culturas_br, y=producao_mt,
                      color=culturas_br,
                      color_discrete_sequence=px.colors.sequential.Greens,
                      template="plotly_dark", labels={"x":"Cultura","y":"Produção (Mt)"},
                      title="Produção em Milhões de Toneladas (2023/24)")
        fig2.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=280, showlegend=False)
        st.plotly_chart(fig2, use_container_width=True)



# ══════════════════════════════════════════════════════════════════
# 🔄 PIPELINE RPA ESPACIAL
# ══════════════════════════════════════════════════════════════════
elif menu == "🔄 Automação Robótica de Processos":
    st.title("🔄 RPA Espacial - Automação Inteligente de Processos")
    st.markdown("Visualize e simule o fluxo completo de automação inspirado na exploração espacial: da coleta de dados à decisão inteligente.")
    st.markdown("---")

    ETAPAS_RPA = [
        {
            "icone": "🛰️", "nome": "1. Coleta Remota",
            "desc": "Satélites e drones capturam imagens multiespectrais, dados climáticos e telemetria de campo de forma contínua e automatizada.",
            "tecnologia": "Sentinel-2 · Landsat · IoT LoRaWAN · MQTT",
            "analogia": "Como sistemas de telemetria de uma espaçonave monitorando todos os sensores em tempo real.",
            "status": "✅ Automático", "cor": "#1f6feb"
        },
        {
            "icone": "⚙️", "nome": "2. Ingestão & Pré-processamento",
            "desc": "Dados brutos são automaticamente ingeridos, corrigidos atmosfericamente, normalizados e indexados em repositórios estruturados.",
            "tecnologia": "Google Earth Engine · Apache Kafka · ETL Pipeline",
            "analogia": "Como o processamento de dados brutos de missão antes de enviá-los ao centro de controle.",
            "status": "✅ Automático", "cor": "#1f6feb"
        },
        {
            "icone": "🧠", "nome": "3. Análise por IA",
            "desc": "Modelos de ML/DL processam os dados: classificam culturas, calculam índices vegetativos, detectam anomalias e estimam produtividade.",
            "tecnologia": "CNN · XGBoost · LSTM · NDVI · EVI · NDWI",
            "analogia": "Como a IA de bordo de um satélite que processa imagens autonomamente sem aguardar comando terrestre.",
            "status": "✅ Automático", "cor": "#40c463"
        },
        {
            "icone": "📄", "nome": "4. Análise de Documentos",
            "desc": "Relatórios técnicos, laudos agronômicos e PDFs são lidos e interpretados automaticamente pela IA, extraindo insights chave.",
            "tecnologia": "RAG · LLM · pdfplumber · LangChain · FAISS",
            "analogia": "Como sistemas de parsing de telemetria que interpretam automaticamente logs de missão.",
            "status": "✅ Automático", "cor": "#40c463"
        },
        {
            "icone": "📊", "nome": "5. Análise de Planilhas",
            "desc": "Dados tabulares de campo (CSVs, XLSXs) são ingeridos, estatisticamente analisados e visualizados com alertas automáticos.",
            "tecnologia": "Pandas · Plotly · NumPy · OpenPyXL",
            "analogia": "Como a análise automática de telemetria estruturada registrada pelos sistemas de bordo.",
            "status": "✅ Automático", "cor": "#40c463"
        },
        {
            "icone": "📋", "nome": "6. Geração de Relatório",
            "desc": "Relatórios executivos são gerados automaticamente com métricas, diagnósticos, alertas e recomendações prontos para decisão.",
            "tecnologia": "Streamlit · Plotly · PDF Export · LLM Summary",
            "analogia": "Como o sumário de missão gerado automaticamente ao final de cada órbita.",
            "status": "⚡ Semi-automático", "cor": "#e3b341"
        },
        {
            "icone": "✅", "nome": "7. Decisão & Ação",
            "desc": "O gestor recebe recomendações prioritizadas e aciona ações corretivas: ajuste de irrigação, aplicação de defensivo, alerta de área desmatada.",
            "tecnologia": "Dashboard · Alertas · API Integração · Notificações",
            "analogia": "Como o centro de controle de missão que recebe os dados processados e decide a próxima manobra.",
            "status": "👤 Humano supervisionado", "cor": "#e3b341"
        },
    ]

    # Visualização visual do pipeline
    st.markdown("### 🗺️ Fluxo de etapas para Pipeline")
    cols_pipe = st.columns(len(ETAPAS_RPA))
    for i, etapa in enumerate(ETAPAS_RPA):
        with cols_pipe[i]:
            st.markdown(f"""
            <div style='background:#1a2744; border-radius:10px; padding:10px 6px;
                        border-top:4px solid {etapa["cor"]}; text-align:center; min-height:110px;'>
              <div style='font-size:1.6em;'>{etapa["icone"]}</div>
              <div style='color:{etapa["cor"]}; font-size:0.72em; font-weight:bold; line-height:1.2;'>{etapa["nome"]}</div>
              <div style='color:#aaa; font-size:0.65em; margin-top:4px;'>{etapa["status"]}</div>
            </div>
            """, unsafe_allow_html=True)
            if i < len(ETAPAS_RPA) - 1:
                pass  # arrow handled by columns layout

    st.markdown("---")

    # Detalhes de cada etapa
    st.markdown("### 🔍 Detalhes do Pipeline — clique em uma etapa")
    sel = st.selectbox("Selecione a etapa para ver detalhes:", [e["nome"] for e in ETAPAS_RPA])
    etapa_info = next(e for e in ETAPAS_RPA if e["nome"] == sel)

    st.markdown(f"""
    <div style='background:#0d2137; border-left:5px solid {etapa_info["cor"]};
                border-radius:12px; padding:20px; margin:10px 0;'>
      <h3 style='color:{etapa_info["cor"]}; margin:0 0 10px 0;'>{etapa_info["icone"]} {etapa_info["nome"]}</h3>
      <p style='color:#e8f5e9;'><b>O que faz:</b> {etapa_info["desc"]}</p>
      <p style='color:#aaa;'><b>🔧 Tecnologias:</b> {etapa_info["tecnologia"]}</p>
      <p style='color:#f0c040;'><b>🚀 Analogia Espacial:</b> {etapa_info["analogia"]}</p>
      <span style='background:{etapa_info["cor"]}33; color:{etapa_info["cor"]};
                   border-radius:6px; padding:4px 10px; font-size:0.85em;'>{etapa_info["status"]}</span>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    # Simulação animada do pipeline
    st.markdown("### ▶️ Simular Analise do Pipeline")
    col_sim1, col_sim2 = st.columns([1, 3])
    with col_sim1:
        if st.button("🚀 Analise Pipeline", use_container_width=True):
            st.session_state["rpa_simulando"] = True

    if st.session_state.get("rpa_simulando"):
        barra = st.progress(0)
        status_txt = st.empty()
        log_area = st.empty()
        logs = []
        for i, etapa in enumerate(ETAPAS_RPA):
            pct = int((i + 1) / len(ETAPAS_RPA) * 100)
            barra.progress(pct)
            status_txt.markdown(f"**Executando:** {etapa['icone']} {etapa['nome']} — {etapa['status']}")
            logs.append(f"✅ [{time.strftime('%H:%M:%S')}] {etapa['nome']} — concluído")
            log_area.code("\n".join(logs), language="bash")
            time.sleep(0.6)
        status_txt.markdown("### 🎯 Pipeline concluído com sucesso!")
        st.session_state["rpa_simulando"] = False
        

    st.markdown("---")
    st.markdown("### 📈 Métricas do Pipeline")
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("⚡ Etapas Automatizadas", "5 de 7", "71%")
    with c2: st.metric("⏱️ Tempo Total", "~4 min", "vs 2h manual")
    with c3: st.metric("🎯 Precisão Média", "91%", "+23% vs humano")
    with c4: st.metric("💰 Redução de Custo", "68%", "vs processo manual")


# ══════════════════════════════════════════════════════════════════
# 📄 ANÁLISE DE DOCUMENTOS IA
# ══════════════════════════════════════════════════════════════════
elif menu == "📄 Pesquisa Doc":
    st.title("📄 Análise de Documentos por IA")
    st.markdown("Faça upload de **PDFs** ou **Word (.docx)** e a IA extrairá insights, resumos e pontos-chave automaticamente.")
    st.markdown("---")

    # ── Verificar bibliotecas disponíveis ─────────────────────────
    libs_ok = {}
    try:
        import pdfplumber
        libs_ok["pdfplumber"] = True
    except ImportError:
        libs_ok["pdfplumber"] = False

    try:
        import fitz  # PyMuPDF
        libs_ok["pymupdf"] = True
    except ImportError:
        libs_ok["pymupdf"] = False

    try:
        import pypdf
        libs_ok["pypdf"] = True
    except ImportError:
        try:
            import PyPDF2
            libs_ok["pypdf"] = True
            libs_ok["pypdf_legacy"] = True
        except ImportError:
            libs_ok["pypdf"] = False
            libs_ok["pypdf_legacy"] = False

    try:
        from docx import Document as _DocxTest
        libs_ok["docx"] = True
    except ImportError:
        libs_ok["docx"] = False

    pdf_disponivel = libs_ok["pdfplumber"] or libs_ok["pymupdf"] or libs_ok.get("pypdf", False)

    # Mostrar status das bibliotecas
    with st.expander("🔧 Status das bibliotecas de leitura"):
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown(f"{'✅' if libs_ok['pdfplumber'] else '❌'} **pdfplumber**")
            if not libs_ok["pdfplumber"]:
                st.caption("pip install pdfplumber")
        with c2:
            st.markdown(f"{'✅' if libs_ok['pymupdf'] else '❌'} **PyMuPDF**")
            if not libs_ok["pymupdf"]:
                st.caption("pip install pymupdf")
        with c3:
            st.markdown(f"{'✅' if libs_ok.get('pypdf') else '❌'} **pypdf** (fallback)")
            if not libs_ok.get("pypdf"):
                st.caption("pip install pypdf")
        with c4:
            st.markdown(f"{'✅' if libs_ok['docx'] else '❌'} **python-docx**")
            if not libs_ok["docx"]:
                st.caption("pip install python-docx")

    if not pdf_disponivel:
        st.warning("⚠️ Nenhuma biblioteca de PDF instalada. Instale uma delas no terminal:\n\n`pip install pypdf` ou `pip install pdfplumber`\n\nO upload de .docx continua funcionando normalmente.")

    with st.spinner("⏳ Inicializando motor de IA..."):
        vectorstore, llm_client, _ = init_rag()

    tipo_doc = st.radio("Tipo de documento:", ["📕 PDF", "📘 Word (.docx)"], horizontal=True)

    # Definir tipos aceitos com base na disponibilidade
    if "PDF" in tipo_doc:
        if pdf_disponivel:
            tipos_aceitos = ["pdf"]
        else:
            st.info("📌 Instale `pypdf` ou `pdfplumber` para habilitar o upload de PDF. Por enquanto, use Word (.docx).")
            tipos_aceitos = []
    else:
        tipos_aceitos = ["docx"]

    uploaded_doc = st.file_uploader(
        "📁 Selecione o documento",
        type=tipos_aceitos if tipos_aceitos else ["pdf", "docx"],
        disabled=(not tipos_aceitos and "PDF" in tipo_doc)
    )

    # ── Funções de extração ────────────────────────────────────────
    def extrair_texto_pdf(file_bytes):
        """Tenta pdfplumber primeiro; fallback para PyMuPDF. Retorna (texto, paginas, metodo, aviso)."""
        texto = ""
        paginas_com_texto = 0
        paginas_total = 0
        metodo = ""
        aviso = ""

        # Tentativa 1: pdfplumber
        if libs_ok["pdfplumber"]:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    paginas_total = len(pdf.pages)
                    limite = min(paginas_total, 40)
                    for i, page in enumerate(pdf.pages[:limite]):
                        t = page.extract_text()
                        if t and t.strip():
                            texto += f"\n--- Página {i+1} ---\n{t.strip()}"
                            paginas_com_texto += 1
                metodo = "pdfplumber"
            except Exception as e:
                aviso = f"pdfplumber falhou: {e}. "

        # Tentativa 2: PyMuPDF (fallback)
        if not texto.strip() and libs_ok["pymupdf"]:
            try:
                import fitz
                doc_fitz = fitz.open(stream=file_bytes, filetype="pdf")
                paginas_total = len(doc_fitz)
                limite = min(paginas_total, 40)
                for i in range(limite):
                    page = doc_fitz[i]
                    t = page.get_text()
                    if t and t.strip():
                        texto += f"\n--- Página {i+1} ---\n{t.strip()}"
                        paginas_com_texto += 1
                doc_fitz.close()
                metodo = "PyMuPDF"
            except Exception as e:
                aviso += f"PyMuPDF falhou: {e}."

        # Tentativa 3: pypdf / PyPDF2 (segundo fallback)
        if not texto.strip() and libs_ok.get("pypdf"):
            try:
                if not libs_ok.get("pypdf_legacy"):
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                else:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                paginas_total = len(reader.pages)
                limite = min(paginas_total, 40)
                for i in range(limite):
                    t = reader.pages[i].extract_text()
                    if t and t.strip():
                        texto += f"\n--- Página {i+1} ---\n{t.strip()}"
                        paginas_com_texto += 1
                metodo = "pypdf"
            except Exception as e:
                aviso += f"pypdf falhou: {e}."

        return texto.strip(), paginas_com_texto, paginas_total, metodo, aviso

    def extrair_texto_docx(file_bytes):
        """Extrai texto de arquivo .docx incluindo tabelas."""
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(file_bytes))
            partes = []

            # Parágrafos normais
            for p in doc.paragraphs:
                if p.text.strip():
                    partes.append(p.text.strip())

            # Texto dentro de tabelas
            for table in doc.tables:
                for row in table.rows:
                    linha = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if linha:
                        partes.append(linha)

            return "\n".join(partes), len(partes), 1, "python-docx", ""
        except Exception as e:
            return "", 0, 0, "", f"Erro ao ler DOCX: {e}"

    def analisar_documento_com_ia(texto, client, tipo_analise):
        prompts = {
            "📝 Resumo Executivo":
                "Faça um resumo executivo claro e objetivo deste documento em até 5 parágrafos. Destaque os pontos mais importantes.",
            "🔑 Pontos-Chave":
                "Liste os 7 pontos-chave mais importantes deste documento em formato de tópicos numerados. Seja específico.",
            "⚠️ Alertas e Riscos":
                "Identifique todos os alertas, riscos, problemas ou pontos críticos mencionados neste documento. Liste cada um.",
            "📊 Dados e Métricas":
                "Extraia TODOS os dados numéricos, métricas, percentuais, datas e indicadores mencionados no documento. Liste-os organizados.",
            "✅ Recomendações":
                "Liste todas as recomendações, sugestões de ação e próximos passos presentes no documento.",
        }
        prompt = prompts.get(tipo_analise, prompts["📝 Resumo Executivo"])

        # Usa até 8000 caracteres para melhor contexto
        texto_para_ia = texto[:8000]

        try:
            resp = client.chat.completions.create(
                model="meta-llama/Llama-3.1-8B-Instruct",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Você é um analista especializado em agricultura, meio ambiente e sensoriamento remoto. "
                            "Responda SEMPRE em português brasileiro, de forma clara, estruturada e detalhada. "
                            "Base suas respostas EXCLUSIVAMENTE no documento fornecido. "
                            "Se o documento não contiver informação suficiente, diga isso claramente."
                        )
                    },
                    {
                        "role": "user",
                        "content": f"{prompt}\n\n===DOCUMENTO===\n{texto_para_ia}\n===FIM DO DOCUMENTO==="
                    }
                ],
                max_tokens=1000,
                temperature=0.2
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            return f"❌ Erro na API de IA: {type(e).__name__}: {e}"

    # ── Interface principal ────────────────────────────────────────
    if uploaded_doc:
        file_bytes = uploaded_doc.read()
        tamanho_kb = len(file_bytes) / 1024

        st.success(f"✅ Arquivo recebido: **{uploaded_doc.name}** — {tamanho_kb:.1f} KB")

        # Extração de texto com diagnóstico completo
        with st.spinner("📖 Extraindo texto do documento... aguarde"):
            if "PDF" in tipo_doc:
                texto_extraido, pags_texto, pags_total, metodo, aviso = extrair_texto_pdf(file_bytes)
            else:
                texto_extraido, pags_texto, pags_total, metodo, aviso = extrair_texto_docx(file_bytes)

        # ── Diagnóstico da extração ────────────────────────────────
        st.markdown("### 🔬 Diagnóstico da Extração")
        col_d1, col_d2, col_d3, col_d4 = st.columns(4)
        with col_d1: st.metric("📄 Total de páginas", pags_total)
        with col_d2: st.metric("✅ Páginas com texto", pags_texto)
        with col_d3: st.metric("📝 Palavras extraídas", f"{len(texto_extraido.split()):,}" if texto_extraido else "0")
        with col_d4: st.metric("🔧 Método usado", metodo if metodo else "—")

        if aviso:
            st.warning(f"⚠️ {aviso}")

        # ── Casos de falha bem explicados ─────────────────────────
        if not texto_extraido or len(texto_extraido.strip()) < 50:
            st.error("❌ **Não foi possível extrair texto deste documento.**")

            st.markdown("#### 🔍 Por que isso acontece?")
            st.markdown("""
            **Causa mais comum — PDF escaneado (imagem):**
            O arquivo é uma foto/digitalização de papel. O PDF não contém texto real,
            apenas uma imagem da página. Nenhuma biblioteca de extração de texto consegue ler isso
            sem OCR (reconhecimento óptico de caracteres).

            **Outras causas possíveis:**
            - PDF protegido por senha
            - PDF gerado por software que não incorporou o texto corretamente
            - Arquivo corrompido

            **Como resolver:**
            """)

            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown("""
                **Opção 1 — Converter o PDF para texto:**
                - Abra o PDF no Adobe Acrobat
                - Vá em Ferramentas → OCR → Reconhecer Texto
                - Salve e faça upload novamente

                **Opção 2 — Usar Google Drive:**
                - Faça upload do PDF no Google Drive
                - Clique com botão direito → Abrir com Google Docs
                - O Google faz OCR automaticamente
                - Baixe como .docx e faça upload aqui
                """)
            with col_s2:
                st.markdown("""
                **Opção 3 — Ferramentas online gratuitas:**
                - https://www.ilovepdf.com/ocr-pdf
                - https://smallpdf.com/pdf-to-word
                - https://pdf2doc.com

                **Opção 4 — Se tiver Python instalado:**
                ```bash
                pip install pytesseract pdf2image
                ```
                (requer Tesseract OCR instalado no sistema)
                """)

            # Mesmo assim, mostrar o que foi possível extrair
            if texto_extraido:
                with st.expander("👁️ Ver o pouco que foi extraído"):
                    st.text(texto_extraido[:500])

        else:
            # ── Extração bem-sucedida ──────────────────────────────
            st.success(f"✅ Texto extraído com sucesso via **{metodo}** — {len(texto_extraido):,} caracteres de {pags_texto} página(s)")

            # Preview do texto extraído
            with st.expander("👁️ Ver texto extraído (clique para expandir)"):
                st.text_area(
                    "Conteúdo extraído:",
                    value=texto_extraido[:3000] + ("\n\n... [texto continua]" if len(texto_extraido) > 3000 else ""),
                    height=300,
                    disabled=True
                )
                st.caption(f"Mostrando os primeiros 3.000 de {len(texto_extraido):,} caracteres extraídos.")

            st.markdown("---")
            st.markdown("### 🧠 Análise por Inteligência Artificial")
            st.markdown("Escolha o que você quer que a IA faça com o documento:")

            tipo_analise = st.selectbox("Tipo de análise:", [
                "📝 Resumo Executivo",
                "🔑 Pontos-Chave",
                "⚠️ Alertas e Riscos",
                "📊 Dados e Métricas",
                "✅ Recomendações",
            ])

            # Descrição do que cada análise faz
            descricoes = {
                "📝 Resumo Executivo": "A IA vai escrever um resumo do documento em até 5 parágrafos.",
                "🔑 Pontos-Chave": "A IA vai listar os 7 pontos mais importantes em tópicos numerados.",
                "⚠️ Alertas e Riscos": "A IA vai identificar tudo que é crítico, problemático ou urgente no documento.",
                "📊 Dados e Métricas": "A IA vai extrair todos os números, percentuais, datas e indicadores.",
                "✅ Recomendações": "A IA vai listar as ações sugeridas ou recomendações presentes no documento.",
            }
            st.info(f"ℹ️ {descricoes[tipo_analise]}")

            if st.button("🚀 Analisar Documento com IA", use_container_width=False):
                with st.spinner(f"🤖 A IA está lendo o documento e gerando: {tipo_analise} — aguarde..."):
                    resultado = analisar_documento_com_ia(texto_extraido, llm_client, tipo_analise)

                if resultado.startswith("❌"):
                    st.error(resultado)
                else:
                    st.markdown(f"""
                    <div style='background:#1a2f1a; border-left:4px solid #40c463;
                                border-radius:10px; padding:20px; margin-top:10px;'>
                      <b style='color:#40c463; font-size:1.1em;'>🛰️ Resultado — {tipo_analise}</b>
                      <hr style='border-color:#2a3f2a; margin:10px 0;'>
                      <span style='color:#e8f5e9; line-height:1.7;'>{resultado.replace(chr(10), '<br>')}</span>
                    </div>
                    """, unsafe_allow_html=True)

            # ── Pergunta livre ─────────────────────────────────────
            st.markdown("---")
            st.markdown("### 💬 Pergunte qualquer coisa sobre o documento")
            st.caption("A IA responderá com base exclusivamente no conteúdo do arquivo que você enviou.")

            pergunta_doc = st.text_input(
                "Sua pergunta:",
                placeholder="Ex: Qual foi a área desmatada? / Quais são as conclusões? / O documento menciona NDVI?"
            )

            if st.button("💬 Perguntar à IA") and pergunta_doc.strip():
                with st.spinner("🤖 Consultando o documento..."):
                    try:
                        texto_trunc = texto_extraido[:8000]
                        resp = llm_client.chat.completions.create(
                            model="meta-llama/Llama-3.1-8B-Instruct",
                            messages=[
                                {
                                    "role": "system",
                                    "content": (
                                        "Você é um analista especializado. Responda em português com base "
                                        "APENAS no documento fornecido. Se a resposta não estiver no documento, "
                                        "diga: 'Essa informação não consta no documento enviado.'"
                                    )
                                },
                                {
                                    "role": "user",
                                    "content": f"===DOCUMENTO===\n{texto_trunc}\n===FIM===\n\nPERGUNTA: {pergunta_doc}"
                                }
                            ],
                            max_tokens=600,
                            temperature=0.2
                        )
                        resposta_ia = resp.choices[0].message.content.strip()
                        st.markdown(f"""
                        <div style='background:#1a2f1a; border-left:4px solid #1f6feb;
                                    border-radius:10px; padding:16px; margin-top:8px;'>
                          <b style='color:#1f6feb;'>🛰️ IA responde:</b><br><br>
                          <span style='color:#e8f5e9; line-height:1.7;'>{resposta_ia.replace(chr(10), '<br>')}</span>
                        </div>
                        """, unsafe_allow_html=True)
                    except Exception as e:
                        st.error(f"❌ Erro na API: {type(e).__name__}: {e}")

    else:
        # ── Tela inicial sem arquivo ───────────────────────────────
        st.info("👆 Selecione o tipo de documento acima e faça o upload para começar.")

        st.markdown("### 💡 Que tipo de documento posso enviar?")
        exemplos_doc = [
            ("📋 Laudo Agronômico (.pdf ou .docx)",
             "A IA extrai recomendações de manejo, alertas fitossanitários e métricas de produtividade"),
            ("🌳 Relatório PRODES/INPE (.pdf)",
             "A IA identifica áreas desmatadas, taxas de supressão e biomas afetados"),
            ("📜 Artigo Científico (.pdf)",
             "A IA resume metodologia, resultados e conclusões do estudo"),
            ("🛰️ Relatório de Missão Espacial (.pdf)",
             "A IA resume dados de telemetria, anomalias e resultados de sensores"),
            ("📄 Qualquer documento Word (.docx)",
             "Relatórios, memorandos, atas, planos de manejo — qualquer texto em .docx"),
        ]
        for titulo, descricao in exemplos_doc:
            st.markdown(f"""
            <div style='background:#1a2744; border-radius:8px; padding:12px; margin:6px 0;
                        border-left:3px solid #40c463;'>
              <b style='color:#40c463;'>{titulo}</b><br>
              <small style='color:#aaa;'>{descricao}</small>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### ⚠️ O que NÃO funciona")
        st.markdown("""
        - **PDF escaneado (foto de papel):** não tem texto real, só imagem — a IA não consegue ler
        - **PDF protegido por senha:** bloqueado para extração
        - **Imagens .jpg/.png:** não são documentos de texto

        **Dica:** Se seu PDF não funcionar, abra-o no Google Drive e salve como .docx.
        """)


# ══════════════════════════════════════════════════════════════════
# 📊 ANÁLISE DE PLANILHAS IA
# ══════════════════════════════════════════════════════════════════
elif menu == "📊 Análise de Planilhas IA":
    st.title("📊 Análise de Planilhas por IA")
    st.markdown("Faça upload de **CSV** ou **XLSX** e a IA analisa automaticamente os dados, gera gráficos e insights.")
    st.markdown("---")

    with st.spinner("⏳ Inicializando motor de IA..."):
        vectorstore, llm_client, _ = init_rag()

    uploaded_sheet = st.file_uploader(
        "📁 Selecione a planilha (CSV ou XLSX)",
        type=["csv", "xlsx", "xls"]
    )

    if uploaded_sheet:
        file_bytes = uploaded_sheet.read()
        nome = uploaded_sheet.name

        with st.spinner("📖 Lendo planilha..."):
            try:
                if nome.endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(file_bytes))
                else:
                    df = pd.read_excel(io.BytesIO(file_bytes))
                erro_leitura = None
            except Exception as e:
                df = None
                erro_leitura = str(e)

        if df is not None:
            st.success(f"✅ Planilha carregada: **{nome}** — {df.shape[0]:,} linhas × {df.shape[1]} colunas")

            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            with col_m1: st.metric("📋 Linhas", f"{df.shape[0]:,}")
            with col_m2: st.metric("📐 Colunas", df.shape[1])
            with col_m3: st.metric("🔢 Numéricas", len(df.select_dtypes(include='number').columns))
            with col_m4: st.metric("📝 Texto", len(df.select_dtypes(exclude='number').columns))

            st.markdown("---")
            tab1, tab2, tab3, tab4 = st.tabs(["👁️ Prévia", "📈 Gráficos", "📊 Estatísticas", "🧠 Insights IA"])

            with tab1:
                st.markdown("#### Primeiras linhas da planilha")
                st.dataframe(df.head(20), use_container_width=True)
                st.caption(f"Exibindo 20 de {df.shape[0]:,} linhas.")

            with tab2:
                st.markdown("#### 📈 Visualização Automática")
                colunas_num = df.select_dtypes(include='number').columns.tolist()
                colunas_cat = df.select_dtypes(exclude='number').columns.tolist()

                if colunas_num:
                    col_x = st.selectbox("Eixo X:", df.columns.tolist(), key="px")
                    col_y = st.selectbox("Eixo Y (numérico):", colunas_num, key="py")
                    tipo_graf = st.radio("Tipo de gráfico:", ["📊 Barras", "📈 Linha", "⚫ Dispersão", "🥧 Pizza"], horizontal=True)

                    if tipo_graf == "📊 Barras":
                        fig = px.bar(df, x=col_x, y=col_y, template="plotly_dark",
                                     color=col_y, color_continuous_scale="Greens")
                    elif tipo_graf == "📈 Linha":
                        fig = px.line(df, x=col_x, y=col_y, template="plotly_dark",
                                      color_discrete_sequence=["#40c463"], markers=True)
                    elif tipo_graf == "⚫ Dispersão":
                        fig = px.scatter(df, x=col_x, y=col_y, template="plotly_dark",
                                         color_discrete_sequence=["#40c463"])
                    else:
                        if colunas_cat:
                            fig = px.pie(df, names=col_x, values=col_y, template="plotly_dark",
                                         hole=0.3, color_discrete_sequence=px.colors.sequential.Greens)
                        else:
                            fig = px.bar(df, x=col_x, y=col_y, template="plotly_dark")

                    fig.update_layout(margin=dict(l=0,r=0,t=30,b=0), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Nenhuma coluna numérica encontrada para visualização.")

            with tab3:
                st.markdown("#### 📊 Estatísticas Descritivas")
                colunas_num = df.select_dtypes(include='number').columns.tolist()
                if colunas_num:
                    st.dataframe(df[colunas_num].describe().round(2), use_container_width=True)

                    # Alertas automáticos
                    st.markdown("#### ⚠️ Alertas Automáticos")
                    nulos = df.isnull().sum()
                    nulos_sig = nulos[nulos > 0]
                    if len(nulos_sig) > 0:
                        for col, qtd in nulos_sig.items():
                            pct = qtd / len(df) * 100
                            if pct > 20:
                                st.error(f"🔴 Coluna **{col}**: {qtd} valores nulos ({pct:.1f}%) — requer atenção!")
                            else:
                                st.warning(f"🟡 Coluna **{col}**: {qtd} valores nulos ({pct:.1f}%)")
                    else:
                        st.success("✅ Nenhum valor nulo encontrado na planilha.")

                    # Mapa de calor de correlação
                    if len(colunas_num) >= 2:
                        st.markdown("#### 🔥 Correlação entre variáveis numéricas")
                        corr = df[colunas_num].corr().round(2)
                        fig_corr = px.imshow(corr, template="plotly_dark",
                                             color_continuous_scale="RdYlGn",
                                             title="Mapa de Correlação")
                        fig_corr.update_layout(height=350, margin=dict(l=0,r=0,t=40,b=0))
                        st.plotly_chart(fig_corr, use_container_width=True)
                else:
                    st.info("Sem colunas numéricas para estatísticas.")

            with tab4:
                st.markdown("#### 🧠 Insights por Inteligência Artificial")

                if st.button("🚀 Gerar Insights com IA", use_container_width=False):
                    resumo_df = f"""
Planilha: {nome}
Dimensões: {df.shape[0]} linhas × {df.shape[1]} colunas
Colunas: {', '.join(df.columns.tolist())}
Tipos: {df.dtypes.to_dict()}
Estatísticas:
{df.describe().round(2).to_string() if not df.select_dtypes(include='number').empty else 'Sem colunas numéricas'}
Primeiras linhas:
{df.head(5).to_string()}
"""
                    with st.spinner("🤖 Analisando dados com IA..."):
                        try:
                            resp = llm_client.chat.completions.create(
                                model="meta-llama/Llama-3.1-8B-Instruct",
                                messages=[
                                    {"role": "system", "content": "Você é um cientista de dados especializado em agricultura e sensoriamento remoto. Analise os dados fornecidos e gere insights práticos em português."},
                                    {"role": "user", "content": f"Analise esta planilha e forneça:\n1. Interpretação dos dados\n2. Padrões ou tendências identificadas\n3. Anomalias ou alertas\n4. Recomendações de ação\n\nDADOS:\n{resumo_df[:3000]}"}
                                ],
                                max_tokens=700, temperature=0.4
                            )
                            resultado_ia = resp.choices[0].message.content.strip()
                            st.markdown(f"""
                            <div style='background:#1a2f1a; border-left:4px solid #40c463;
                                        border-radius:10px; padding:18px;'>
                              <b style='color:#40c463;'>🛰️ Análise da IA</b><br><br>
                              <span style='color:#e8f5e9;'>{resultado_ia}</span>
                            </div>
                            """, unsafe_allow_html=True)
                        except Exception as e:
                            st.error(f"Erro na análise: {e}")

                st.markdown("---")
                st.markdown("#### 💬 Pergunte sobre os dados")
                pergunta_plan = st.text_input("Sua pergunta sobre a planilha:", key="q_plan")
                if st.button("Perguntar sobre dados") and pergunta_plan.strip():
                    with st.spinner("🤖 Consultando IA..."):
                        try:
                            dados_resumo = f"Colunas: {list(df.columns)}\nEstatísticas:\n{df.describe().round(2).to_string()}\nAmostra:\n{df.head(8).to_string()}"
                            resp = llm_client.chat.completions.create(
                                model="meta-llama/Llama-3.1-8B-Instruct",
                                messages=[
                                    {"role": "system", "content": "Analista de dados agrícolas. Responda com base nos dados fornecidos, em português."},
                                    {"role": "user", "content": f"DADOS:\n{dados_resumo[:3000]}\n\nPERGUNTA: {pergunta_plan}"}
                                ],
                                max_tokens=500, temperature=0.3
                            )
                            st.markdown(f"""
                            <div class='chat-bot'>🛰️ <b>IA:</b> {resp.choices[0].message.content.strip()}</div>
                            """, unsafe_allow_html=True)
                        except Exception as e:
                            st.error(f"Erro: {e}")
        else:
            st.error(f"❌ Erro ao ler planilha: {erro_leitura}")
    else:
        st.info("👆 Faça upload de um CSV ou XLSX para iniciar a análise automática.")

        st.markdown("### 💡 Exemplos de planilhas para análise")
        exemplos_plan = [
            ("🌱 Dados de NDVI por talhão", "Área, NDVI_Jan, NDVI_Fev... → identifica tendências sazonais"),
            ("💧 Consumo hídrico por cultura", "Cultura, Lâmina_mm, Precipitação → otimiza irrigação"),
            ("🌳 Áreas de vegetação nativa (CAR)", "Propriedade, Área_ha, Lat, Lon → controle de preservação"),
            ("🛰️ Telemetria de sensores IoT", "Timestamp, Temp, Umidade, NDVI → séries temporais"),
        ]
        for titulo, descricao in exemplos_plan:
            st.markdown(f"""
            <div style='background:#1a2744; border-radius:8px; padding:12px; margin:6px 0;
                        border-left:3px solid #1f6feb;'>
              <b style='color:#1f6feb;'>{titulo}</b><br>
              <small style='color:#aaa;'>{descricao}</small>
            </div>
            """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# 🌐 WEB SCRAPING & XPATH
# ══════════════════════════════════════════════════════════════════
elif menu == "🌐 Módulo de coleta automatizada de dados":
    st.title("🌐 Módulo de coleta automatizada de dados")
    st.markdown("---")

    aba1, aba2, aba3 = st.tabs([
        "🌐 Scraping + XPath HTML",
        "🗺️ Parser XML/KML Geográfico",
        "🤖 Selenium RPA Demo"
    ])

    # ── ABA 1: Scraping + XPath HTML ──────────────────────────────
    with aba1:
        st.markdown("### 🌐 Web Scraping com requests + XPath (lxml)")
        st.markdown("""
        Raspa dados de portais públicos e usa **XPath** para navegar no HTML como se fosse XML.
        XPath é uma linguagem de caminho que permite selecionar nós de um documento XML/HTML
        com expressões precisas como `//div[@class='titulo']` ou `//table/tr/td[2]`.
        """)

        st.markdown("#### 🔑 Sintaxe XPath — Referência Rápida")
        df_xpath = pd.DataFrame({
            "Expressão XPath": [
                "//tag", "//tag[@atributo='valor']", "//tag/filho",
                "//tag[1]", "//tag/text()", "//tag/@atributo",
                "//*[contains(@class,'nome')]", "//tag[last()]"
            ],
            "O que seleciona": [
                "Todos os elementos <tag> no documento",
                "Tags com atributo específico",
                "Filhos diretos de tag",
                "Primeiro elemento <tag>",
                "Texto interno da tag",
                "Valor do atributo",
                "Tags cuja classe contém 'nome'",
                "Último elemento da lista"
            ],
            "Exemplo real": [
                "//h1", "//a[@href]", "//table/tr",
                "//li[1]", "//p/text()", "//img/@src",
                "//*[contains(@class,'ndvi')]", "//tr[last()]"
            ]
        })
        st.dataframe(df_xpath, use_container_width=True, hide_index=True)

        st.markdown("---")
        st.markdown("#### 🚀 Executar Scraping ao Vivo")

        FONTES_DISPONIVEIS = {
            "📡 INPE — Queimadas BR": {
                "url": "https://queimadas.dgi.inpe.br/queimadas/bdqueimadas",
                "desc": "Portal de monitoramento de focos de queimadas do INPE",
                "xpath_exemplo": "//title/text()",
                "modo": "requests"
            },
            "🌿 Wikipedia — Amazônia": {
                "url": "https://pt.wikipedia.org/wiki/Amaz%C3%B4nia",
                "desc": "Artigo Wikipedia sobre a Amazônia — extração de parágrafos e dados",
                "xpath_exemplo": "//div[@class='mw-parser-output']/p",
                "modo": "requests"
            },
            "🛰️ Wikipedia — Sentinel-2": {
                "url": "https://en.wikipedia.org/wiki/Sentinel-2",
                "desc": "Informações técnicas do satélite Sentinel-2",
                "xpath_exemplo": "//div[@class='mw-parser-output']/p",
                "modo": "requests"
            },
            "🌱 Embrapa — Notícias": {
                "url": "https://www.embrapa.br/busca-de-noticias/-/noticia/lista/noticias",
                "desc": "Feed de notícias da Embrapa sobre agricultura e pesquisa",
                "xpath_exemplo": "//h3/a/text()",
                "modo": "requests"
            },
        }

        fonte_sel = st.selectbox("Escolha a fonte para raspar:", list(FONTES_DISPONIVEIS.keys()))
        info_fonte = FONTES_DISPONIVEIS[fonte_sel]

        st.info(f"**URL:** `{info_fonte['url']}`\n\n{info_fonte['desc']}")

        xpath_custom = st.text_input(
            "XPath para extração (edite ou use o padrão):",
            value=info_fonte["xpath_exemplo"],
            help="Digite uma expressão XPath para selecionar elementos do HTML"
        )

        col_b1, col_b2 = st.columns([1, 4])
        with col_b1:
            executar_scraping = st.button("🚀 Executar Scraping", use_container_width=True)

        if executar_scraping:
            with st.spinner(f"🌐 Raspando {info_fonte['url']}..."):
                try:
                    headers = {
                        "User-Agent": "Mozilla/5.0 (compatible; AgroSpaceBot/1.0; research)"
                    }
                    response = req_lib.get(info_fonte["url"], headers=headers, timeout=12)
                    response.encoding = response.apparent_encoding
                    html_content = response.text
                    status = response.status_code

                    st.success(f"✅ Conexão estabelecida — HTTP {status} | {len(html_content):,} bytes recebidos")

                    # Parse com lxml + XPath
                    from lxml import html as lxml_html
                    tree = lxml_html.fromstring(html_content.encode("utf-8", errors="replace"))
                    elementos = tree.xpath(xpath_custom)

                    st.markdown(f"#### 📌 Resultado XPath: `{xpath_custom}`")
                    st.markdown(f"**{len(elementos)} elementos encontrados**")

                    if elementos:
                        resultados = []
                        for el in elementos[:15]:
                            if isinstance(el, str):
                                texto = el.strip()
                            else:
                                try:
                                    texto = el.text_content().strip()
                                except:
                                    texto = str(el)
                            if texto and len(texto) > 5:
                                resultados.append(texto[:300])

                        if resultados:
                            for i, r in enumerate(resultados[:10], 1):
                                st.markdown(f"""
                                <div style='background:#1a2744; border-left:3px solid #40c463;
                                            border-radius:8px; padding:10px; margin:4px 0;'>
                                  <small style='color:#aaa;'>#{i}</small><br>
                                  <span style='color:#e8f5e9;'>{r}</span>
                                </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.warning("XPath encontrou elementos mas sem texto legível. Tente `//p/text()` ou `//h2/text()`.")
                    else:
                        st.warning(f"Nenhum elemento encontrado com XPath `{xpath_custom}`. Tente outro seletor.")

                    # Preview do HTML bruto
                    with st.expander("📄 Ver HTML bruto (primeiros 2.000 caracteres)"):
                        st.code(html_content[:2000], language="html")

                    # BeautifulSoup complementar
                    with st.expander("🍲 Análise complementar com BeautifulSoup"):
                        soup = BeautifulSoup(html_content, "html.parser")
                        titulo = soup.title.string if soup.title else "Sem título"
                        links = [(a.get_text(strip=True), a.get("href","")) for a in soup.find_all("a", href=True)[:10]]
                        st.markdown(f"**Título da página:** {titulo}")
                        st.markdown(f"**Total de links encontrados:** {len(soup.find_all('a', href=True))}")
                        st.markdown("**Primeiros 10 links:**")
                        df_links = pd.DataFrame(links, columns=["Texto", "URL"])
                        st.dataframe(df_links, use_container_width=True, hide_index=True)

                except req_lib.exceptions.Timeout:
                    st.error("⏱️ Timeout — o servidor demorou mais de 12s para responder.")
                except req_lib.exceptions.ConnectionError:
                    st.error("❌ Erro de conexão — verifique se a URL está acessível.")
                except Exception as e:
                    st.error(f"❌ Erro: {type(e).__name__}: {e}")

    # ── ABA 2: Parser XML/KML Geográfico ──────────────────────────
    with aba2:
        st.markdown("### 🗺️ Parser XML/KML Geográfico com XPath + lxml")
        st.markdown("""
        Arquivos **KML** (Google Earth) e **GML/XML** geográficos são documentos XML.
        O XPath permite navegar neles com precisão para extrair coordenadas, nomes de áreas,
        atributos ambientais e delimitar polígonos de preservação ou supressão.
        """)

        st.markdown("#### 📂 Modo de entrada")
        modo_xml = st.radio("Fonte do XML/KML:", [
            "📝 Usar exemplo embutido (Área Preservada)",
            "📁 Upload de arquivo XML/KML"
        ], horizontal=True)

        KML_EXEMPLO = '''<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>Mapeamento de Vegetação — AgroSpace</name>
    <description>Áreas monitoradas via sensoriamento remoto — Sentinel-2</description>

    <Placemark>
      <name>Área Preservada 1 — Mata Atlântica</name>
      <description>Vegetação nativa — NDVI médio 0.82</description>
      <ExtendedData>
        <Data name="tipo"><value>preservacao</value></Data>
        <Data name="ndvi_medio"><value>0.82</value></Data>
        <Data name="area_ha"><value>1250.5</value></Data>
        <Data name="bioma"><value>Mata Atlântica</value></Data>
        <Data name="status"><value>Protegida — APP</value></Data>
      </ExtendedData>
      <Polygon>
        <outerBoundaryIs><LinearRing>
          <coordinates>-47.12,-23.45,0 -47.10,-23.45,0 -47.10,-23.47,0 -47.12,-23.47,0 -47.12,-23.45,0</coordinates>
        </LinearRing></outerBoundaryIs>
      </Polygon>
    </Placemark>

    <Placemark>
      <name>Área Suprimida 1 — Conversão Agropecuária</name>
      <description>Desmatamento detectado — NDVI médio 0.18</description>
      <ExtendedData>
        <Data name="tipo"><value>supressao</value></Data>
        <Data name="ndvi_medio"><value>0.18</value></Data>
        <Data name="area_ha"><value>340.2</value></Data>
        <Data name="bioma"><value>Cerrado</value></Data>
        <Data name="status"><value>ALERTA — Supressão ilegal detectada</value></Data>
      </ExtendedData>
      <Polygon>
        <outerBoundaryIs><LinearRing>
          <coordinates>-48.05,-22.10,0 -48.02,-22.10,0 -48.02,-22.13,0 -48.05,-22.13,0 -48.05,-22.10,0</coordinates>
        </LinearRing></outerBoundaryIs>
      </Polygon>
    </Placemark>

    <Placemark>
      <name>Área Preservada 2 — Cerrado Nativo</name>
      <description>Reserva Legal — NDVI médio 0.65</description>
      <ExtendedData>
        <Data name="tipo"><value>preservacao</value></Data>
        <Data name="ndvi_medio"><value>0.65</value></Data>
        <Data name="area_ha"><value>820.0</value></Data>
        <Data name="bioma"><value>Cerrado</value></Data>
        <Data name="status"><value>Reserva Legal — Código Florestal</value></Data>
      </ExtendedData>
      <Polygon>
        <outerBoundaryIs><LinearRing>
          <coordinates>-46.90,-23.30,0 -46.87,-23.30,0 -46.87,-23.33,0 -46.90,-23.33,0 -46.90,-23.30,0</coordinates>
        </LinearRing></outerBoundaryIs>
      </Polygon>
    </Placemark>

    <Placemark>
      <name>Área Suprimida 2 — Queimada Recente</name>
      <description>Área queimada — NDVI médio 0.05</description>
      <ExtendedData>
        <Data name="tipo"><value>supressao</value></Data>
        <Data name="ndvi_medio"><value>0.05</value></Data>
        <Data name="area_ha"><value>175.8</value></Data>
        <Data name="bioma"><value>Amazônia</value></Data>
        <Data name="status"><value>ALERTA CRÍTICO — Queimada ativa</value></Data>
      </ExtendedData>
      <Polygon>
        <outerBoundaryIs><LinearRing>
          <coordinates>-55.20,-10.15,0 -55.17,-10.15,0 -55.17,-10.18,0 -55.20,-10.18,0 -55.20,-10.15,0</coordinates>
        </LinearRing></outerBoundaryIs>
      </Polygon>
    </Placemark>

  </Document>
</kml>'''

        xml_content = None
        if modo_xml == "📝 Usar exemplo embutido (Área Preservada)":
            xml_content = KML_EXEMPLO
            st.code(KML_EXEMPLO[:800] + "\n... (truncado)", language="xml")
        else:
            arq_xml = st.file_uploader("Upload do arquivo XML/KML", type=["xml","kml","gml"])
            if arq_xml:
                xml_content = arq_xml.read().decode("utf-8", errors="replace")
                st.success(f"✅ Arquivo carregado: {arq_xml.name}")
                with st.expander("Ver conteúdo"):
                    st.code(xml_content[:1000], language="xml")

        if xml_content:
            st.markdown("---")
            st.markdown("#### 🔍 Consultas XPath no documento")

            CONSULTAS_PRONTAS = {
                "📋 Todos os nomes de áreas": ".//{*}name/text()",
                "🌿 Apenas áreas PRESERVADAS": ".//{*}Data[@name='tipo'][{*}value='preservacao']/../../{*}name/text()",
                "⚠️ Apenas áreas SUPRIMIDAS": ".//{*}Data[@name='tipo'][{*}value='supressao']/../../{*}name/text()",
                "📊 Todos os valores de NDVI": ".//{*}Data[@name='ndvi_medio']/{*}value/text()",
                "🏞️ Todas as áreas em hectares": ".//{*}Data[@name='area_ha']/{*}value/text()",
                "🌍 Biomas presentes": ".//{*}Data[@name='bioma']/{*}value/text()",
                "🚨 Status / alertas": ".//{*}Data[@name='status']/{*}value/text()",
                "📍 Coordenadas dos polígonos": ".//{*}coordinates/text()",
            }

            consulta_sel = st.selectbox("Consulta XPath predefinida:", list(CONSULTAS_PRONTAS.keys()))
            xpath_xml = st.text_input("Expressão XPath (editável):", value=CONSULTAS_PRONTAS[consulta_sel])

            if st.button("🔍 Executar XPath no XML", use_container_width=False):
                try:
                    root = etree.fromstring(xml_content.encode("utf-8"))
                    resultados = root.xpath(xpath_xml)

                    st.markdown(f"**{len(resultados)} resultado(s) encontrado(s):**")
                    if resultados:
                        for i, r in enumerate(resultados, 1):
                            texto = r.strip() if isinstance(r, str) else etree.tostring(r, encoding="unicode")
                            is_num = texto.replace(".","").lstrip("-").isdigit()
                            if "preserva" in texto.lower() or (is_num and float(texto) >= 0.5):
                                cor = "#27ae60"
                            elif "supressao" in texto.lower() or "alerta" in texto.lower():
                                cor = "#e74c3c"
                            else:
                                cor = "#40c463"
                            st.markdown(f"""
                            <div style='background:#1a2744; border-left:3px solid {cor};
                                        padding:8px 12px; margin:3px 0; border-radius:6px;'>
                              <small style='color:#aaa;'>#{i}</small>
                              <span style='color:#e8f5e9; margin-left:8px;'>{texto[:400]}</span>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.warning("Nenhum resultado. Tente outra expressão XPath.")
                except etree.XPathEvalError as e:
                    st.error(f"❌ Erro XPath: {e}")
                except Exception as e:
                    st.error(f"❌ Erro ao parsear XML: {e}")

            # Dashboard automático a partir do KML
            st.markdown("---")
            st.markdown("#### 📊 Dashboard Automático — Preservação vs Supressão")
            try:
                root = etree.fromstring(xml_content.encode("utf-8"))
                nomes    = root.xpath(".//{*}name/text()")
                tipos    = root.xpath(".//{*}Data[@name='tipo']/{*}value/text()")
                ndvis    = root.xpath(".//{*}Data[@name='ndvi_medio']/{*}value/text()")
                areas    = root.xpath(".//{*}Data[@name='area_ha']/{*}value/text()")
                biomas   = root.xpath(".//{*}Data[@name='bioma']/{*}value/text()")
                status_l = root.xpath(".//{*}Data[@name='status']/{*}value/text()")

                # Filtra o nome do Document
                doc_names = root.xpath(".//{*}Document/{*}name/text()")
                doc_name = doc_names[0] if doc_names else ""
                nomes_areas = [n for n in nomes if n != doc_name]

                if tipos and ndvis and areas:
                    df_kml = pd.DataFrame({
                        "Área": nomes_areas[:len(tipos)],
                        "Tipo": tipos,
                        "NDVI Médio": [float(x) for x in ndvis],
                        "Área (ha)": [float(x) for x in areas],
                        "Bioma": biomas,
                        "Status": status_l,
                    })

                    col_k1, col_k2, col_k3 = st.columns(3)
                    preservadas = df_kml[df_kml["Tipo"]=="preservacao"]
                    suprimidas  = df_kml[df_kml["Tipo"]=="supressao"]
                    with col_k1: st.metric("🌿 Áreas Preservadas", len(preservadas), f"{preservadas['Área (ha)'].sum():.0f} ha")
                    with col_k2: st.metric("⚠️ Áreas Suprimidas",  len(suprimidas),  f"{suprimidas['Área (ha)'].sum():.0f} ha")
                    with col_k3: st.metric("📊 NDVI Médio Geral", f"{df_kml['NDVI Médio'].mean():.2f}", "")

                    st.dataframe(df_kml, use_container_width=True, hide_index=True)

                    col_g1, col_g2 = st.columns(2)
                    with col_g1:
                        fig_tipo = px.pie(
                            df_kml, names="Tipo", values="Área (ha)",
                            color="Tipo",
                            color_discrete_map={"preservacao":"#27ae60","supressao":"#e74c3c"},
                            template="plotly_dark", title="Distribuição por Tipo (ha)", hole=0.35
                        )
                        fig_tipo.update_layout(height=300, margin=dict(l=0,r=0,t=40,b=0))
                        st.plotly_chart(fig_tipo, use_container_width=True)
                    with col_g2:
                        fig_ndvi = px.bar(
                            df_kml, x="Área", y="NDVI Médio",
                            color="Tipo",
                            color_discrete_map={"preservacao":"#27ae60","supressao":"#e74c3c"},
                            template="plotly_dark", title="NDVI por Área"
                        )
                        fig_ndvi.add_hline(y=0.5, line_dash="dash", line_color="yellow",
                                           annotation_text="Limiar saudável (0.5)")
                        fig_ndvi.update_layout(height=300, margin=dict(l=0,r=0,t=40,b=0))
                        st.plotly_chart(fig_ndvi, use_container_width=True)
            except Exception as e:
                st.warning(f"Dashboard não gerado: {e}")

    # ── ABA 3: Selenium RPA Demo ───────────────────────────────────
    with aba3:
        st.markdown("### 🤖 Selenium — Automação de Navegador (RPA)")
        st.markdown("""
        **Selenium** é a principal ferramenta de RPA para automatizar navegadores web.
        Combinado com **XPath**, permite localizar qualquer elemento em uma página e
        interagir com ele: clicar, preencher formulários, fazer download, capturar dados.
        """)

        st.info("ℹ️ **Nota:** Selenium requer Chrome/Firefox instalado localmente. "
                "Esta aba demonstra o código e simula a execução — ideal para uso local ou em servidores com navegador.")

        st.markdown("#### 🏗️ Arquitetura RPA com Selenium + XPath")
        col_arq1, col_arq2 = st.columns(2)
        with col_arq1:
            st.markdown("""
            **Componentes necessários:**
            - `selenium` — controle do navegador
            - `webdriver-manager` — instala o driver automaticamente
            - `lxml` — XPath para parsing pós-scraping
            - `ChromeDriver` ou `GeckoDriver`

            **Instalação:**
            ```bash
            pip install selenium webdriver-manager lxml
            ```
            """)
        with col_arq2:
            st.markdown("""
            **Fluxo RPA típico:**
            1. 🚀 Abrir navegador headless
            2. 🌐 Navegar até a URL alvo
            3. ⏳ Aguardar carregamento (WebDriverWait)
            4. 🔍 Localizar elemento com XPath
            5. 🖱️ Interagir (clicar, digitar, scroll)
            6. 📥 Extrair dado / fazer download
            7. 🔄 Repetir para próxima página
            8. ✅ Fechar navegador
            """)

        st.markdown("---")
        st.markdown("#### 💻 Exemplos de Código Selenium + XPath")

        exemplos_sel = {
            "🌿 Scraping INPE/TerraBrasilis": '''from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options

# Configuração headless (sem janela visual)
options = Options()
options.add_argument("--headless")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

driver = webdriver.Chrome(options=options)

try:
    # Navegar ao portal TerraBrasilis
    driver.get("https://terrabrasilis.dpi.inpe.br/app/dashboard/deforestation/biomes/legal_amazon/increments")

    # Aguardar tabela de dados carregar (até 15 segundos)
    wait = WebDriverWait(driver, 15)
    tabela = wait.until(
        EC.presence_of_element_located((By.XPATH, "//table[@id='tbl-increments']"))
    )

    # XPath para extrair linhas da tabela
    linhas = driver.find_elements(By.XPATH, "//table[@id='tbl-increments']//tr")
    for linha in linhas[1:]:  # pular cabeçalho
        colunas = linha.find_elements(By.TAG_NAME, "td")
        dados = [c.text for c in colunas]
        print(dados)  # ex: ['2024', 'Amazônia', '6518 km²']

finally:
    driver.quit()
''',
            "🔍 Localizar elementos por XPath": '''from selenium import webdriver
from selenium.webdriver.common.by import By

driver = webdriver.Chrome()
driver.get("https://www.embrapa.br/busca-de-noticias")

# ── Diferentes formas de usar XPath ──────────────────────────

# 1. Por tag simples
titulos = driver.find_elements(By.XPATH, "//h3")

# 2. Por atributo específico
links_noticias = driver.find_elements(By.XPATH, "//a[@class='title']")

# 3. Texto que contém palavra-chave
agro_links = driver.find_elements(
    By.XPATH, "//a[contains(text(),'satélite')]"
)

# 4. Posição no DOM
primeiro_resultado = driver.find_element(
    By.XPATH, "(//div[@class='result-item'])[1]"
)

# 5. Pai/filho/irmão
subtitulo = driver.find_element(
    By.XPATH, "//h2[@id='destaque']/following-sibling::p[1]"
)

# 6. Condição múltipla
btn_download = driver.find_element(
    By.XPATH, "//a[@href and contains(@href,'.csv') and @data-type='download']"
)

print(f"Encontrados {len(titulos)} títulos")
driver.quit()
''',
            "📥 Download automático de dados": '''from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
import os, time

# Configurar pasta de download automático
pasta_download = "/home/usuario/dados_inpe"
os.makedirs(pasta_download, exist_ok=True)

options = Options()
options.add_argument("--headless")
options.add_experimental_option("prefs", {
    "download.default_directory": pasta_download,
    "download.prompt_for_download": False,
})

driver = webdriver.Chrome(options=options)

try:
    # Acessar portal com dados de desmatamento
    driver.get("https://terrabrasilis.dpi.inpe.br/downloads/")

    # XPath para encontrar link de download do CSV do PRODES
    btn_csv = driver.find_element(
        By.XPATH,
        "//a[contains(@href,'prodes') and contains(@href,'.csv')]"
    )

    btn_csv.click()  # Inicia download automaticamente
    time.sleep(5)    # Aguardar conclusão

    # Listar arquivos baixados
    arquivos = os.listdir(pasta_download)
    print(f"Arquivos baixados: {arquivos}")

finally:
    driver.quit()
''',
            "🔄 Loop de páginas (Paginação)": '''from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pandas as pd

driver = webdriver.Chrome()
driver.get("https://queimadas.dgi.inpe.br/queimadas/bdqueimadas")

todos_dados = []
pagina = 1

while pagina <= 10:  # máximo 10 páginas
    print(f"Raspando página {pagina}...")

    # XPath para extrair dados da tabela atual
    try:
        linhas = driver.find_elements(
            By.XPATH, "//table[@id='resultado']//tbody/tr"
        )
        for linha in linhas:
            cols = linha.find_elements(By.TAG_NAME, "td")
            if cols:
                todos_dados.append({
                    "data": cols[0].text,
                    "municipio": cols[1].text,
                    "estado": cols[2].text,
                    "focos": cols[3].text,
                })
    except:
        break

    # XPath para botão "Próxima página"
    try:
        btn_proximo = driver.find_element(
            By.XPATH, "//a[@id='resultado_next' and not(@class='disabled')]"
        )
        btn_proximo.click()
        WebDriverWait(driver, 10).until(
            EC.staleness_of(linhas[0])  # aguarda tabela recarregar
        )
        pagina += 1
    except:
        print("Última página alcançada.")
        break

df = pd.DataFrame(todos_dados)
df.to_csv("focos_queimadas.csv", index=False)
print(f"Total: {len(df)} registros salvos.")
driver.quit()
'''
        }

        exemplo_sel = st.selectbox("Escolha o exemplo:", list(exemplos_sel.keys()))
        st.code(exemplos_sel[exemplo_sel], language="python")

        st.markdown("---")
        st.markdown("#### 🎮 Simulador de Execução RPA")
        st.markdown("Simule o que aconteceria ao executar o RPA Selenium em um portal ambiental:")

        url_rpa = st.text_input("URL alvo para simular:", value="https://terrabrasilis.dpi.inpe.br")
        xpath_rpa = st.text_input("XPath do elemento alvo:", value="//table[@id='tbl-increments']//tr")
        acao_rpa = st.selectbox("Ação RPA:", ["🔍 Extrair dados (find_elements)", "🖱️ Clicar (click)", "⌨️ Digitar (send_keys)", "📥 Download (click em link)"])

        if st.button("▶️ Simular Execução RPA"):
            etapas_rpa = [
                ("🚀", "Iniciando ChromeDriver headless...", 0.5),
                ("🌐", f"Navegando para: {url_rpa}", 0.8),
                ("⏳", "Aguardando carregamento da página (WebDriverWait)...", 1.0),
                ("🔍", f"Aplicando XPath: `{xpath_rpa}`", 0.6),
                ("⚡", f"Executando ação: {acao_rpa}", 0.7),
                ("📦", "Coletando resultados e estruturando dados...", 0.5),
                ("✅", "Encerrando navegador (driver.quit())", 0.3),
            ]
            log_sim = st.empty()
            prog_sim = st.progress(0)
            logs_sim = []
            for i, (ico, msg, delay) in enumerate(etapas_rpa):
                time.sleep(delay)
                logs_sim.append(f"{ico} [{time.strftime('%H:%M:%S')}] {msg}")
                log_sim.code("\n".join(logs_sim), language="bash")
                prog_sim.progress(int((i+1)/len(etapas_rpa)*100))

            st.success("✅ **Simulação RPA concluída com sucesso!**")

            # Resultado simulado
            dados_sim = {
                "Área": ["Amazônia Legal","Cerrado","Mata Atlântica","Pantanal","Caatinga"],
                "Desmatamento 2024 (km²)": [6518, 3241, 112, 45, 678],
                "Variação (%)": [-28.2, -15.4, -8.1, -3.2, -19.7],
                "Fonte XPath": [xpath_rpa]*5
            }
            df_sim = pd.DataFrame(dados_sim)
            st.markdown("#### 📊 Dados simulados extraídos pelo RPA:")
            st.dataframe(df_sim, use_container_width=True, hide_index=True)
            fig_sim = px.bar(df_sim, x="Área", y="Desmatamento 2024 (km²)",
                             color="Desmatamento 2024 (km²)", color_continuous_scale="Reds",
                             template="plotly_dark", title="Dados de Desmatamento extraídos via RPA Selenium")
            fig_sim.update_layout(height=300, margin=dict(l=0,r=0,t=40,b=0))
            st.plotly_chart(fig_sim, use_container_width=True)

        st.markdown("---")
        st.markdown("#### 📦 requirements.txt para Selenium")
        st.code("""selenium>=4.18.0
webdriver-manager>=4.0.1
lxml>=5.1.0
beautifulsoup4>=4.12.0
requests>=2.31.0""", language="text")

        st.markdown("#### 🖥️ Instalação do ChromeDriver (Ubuntu/Servidor)")
        st.code("""# Instalar Google Chrome
wget -q -O - https://dl-ssl.google.com/linux/linux_signing_key.pub | apt-key add -
echo "deb [arch=amd64] http://dl.google.com/linux/chrome/deb/ stable main" > /etc/apt/sources.list.d/google.list
apt-get update && apt-get install -y google-chrome-stable

# Instalar ChromeDriver automaticamente via webdriver-manager
pip install webdriver-manager
# No código Python:
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
""", language="bash")


# ══════════════════════════════════════════════════════════════════
# 💬 CHAT RAG
# ══════════════════════════════════════════════════════════════════
elif menu == "💬 Chat AgroSpace":
    st.title("💬 Chat AgroSpace — Assistente de Agricultura Espacial")
    st.markdown("Faça perguntas sobre **agricultura inteligente, satélites, drones, IoT e geotecnologias**.")
    st.markdown("---")

    with st.spinner("⏳ Inicializando RAG (apenas na primeira vez)..."):
        vectorstore, llm_client, _ = init_rag()

    if "mensagens" not in st.session_state:
        st.session_state.mensagens = []

    exemplos = [
        "O que é NDVI e como é calculado?",
        "Quais satélites são usados na agricultura de precisão?",
        "Como funciona a irrigação inteligente com dados satelitais?",
        "O que é LoRaWAN e qual sua aplicação no agro?",
        "Como o blockchain é usado na rastreabilidade de commodities?",
        "Qual é o futuro do mercado de créditos de carbono no Brasil?"
    ]

    st.markdown("**💡 Perguntas de exemplo:**")
    cols_ex = st.columns(3)
    for i, ex in enumerate(exemplos):
        with cols_ex[i % 3]:
            if st.button(ex, key=f"ex_{i}"):
                st.session_state.pergunta_rapida = ex

    st.markdown("---")

    for msg in st.session_state.mensagens:
        if msg["role"] == "user":
            st.markdown(f'<div class="chat-user">👤 <b>Você:</b> {msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-bot">🛰️ <b>Assistente:</b> {msg["content"]}</div>', unsafe_allow_html=True)
            if "fontes" in msg:
                st.caption("📚 Fontes: " + " · ".join(msg["fontes"]))

    pergunta_default = st.session_state.pop("pergunta_rapida", "") if "pergunta_rapida" in st.session_state else ""
    pergunta = st.text_input("✏️ Digite sua pergunta:", value=pergunta_default, key="input_chat")

    col_btn1, col_btn2 = st.columns([1,5])
    with col_btn1:
        enviar = st.button("Enviar 🚀")
    with col_btn2:
        if st.button("🗑️ Limpar conversa"):
            st.session_state.mensagens = []
            st.rerun()

    if enviar and pergunta.strip():
        st.session_state.mensagens.append({"role":"user","content":pergunta})
        with st.spinner("🤖 Gerando resposta..."):
            try:
                resposta, fontes = gerar_resposta_rag(pergunta, vectorstore, llm_client)
                st.session_state.mensagens.append({"role":"assistant","content":resposta,"fontes":fontes})
            except Exception as e:
                st.session_state.mensagens.append({"role":"assistant","content":f"Erro: {str(e)}","fontes":[]})
        st.rerun()

# ══════════════════════════════════════════════════════════════════
# 🖼️ ANÁLISE DE IMAGEM NDVI
# ══════════════════════════════════════════════════════════════════
elif menu == "🖼️ Análise de Imagem NDVI":
    st.title("🖼️ Análise de Imagem NDVI")
    st.markdown("Faça upload de uma imagem agrícola para análise automática de vegetação e diagnóstico do solo.")
    st.markdown("---")

    modo = st.radio("Escolha o modo de análise:",
                    ["📂 Upload de imagem real", "🔬 Simulação sintética (sem upload)"],
                    horizontal=True)

    def calcular_ndvi_e_diagnostico(RED, NIR):
        NDVI = np.clip((NIR - RED) / (NIR + RED + 1e-8), -1, 1)
        aptidao = np.zeros_like(NDVI, dtype=int)
        aptidao[NDVI >= 0.5]                       = 3
        aptidao[(NDVI >= 0.3) & (NDVI < 0.5)]     = 2
        aptidao[(NDVI >= 0.1) & (NDVI < 0.3)]     = 1
        pct = {
            "✅ Apto (NDVI ≥ 0.5)":        round((NDVI>=0.5).mean()*100,1),
            "🟡 Moderado (0.3–0.5)":       round(((NDVI>=0.3)&(NDVI<0.5)).mean()*100,1),
            "🟠 Estressado (0.1–0.3)":     round(((NDVI>=0.1)&(NDVI<0.3)).mean()*100,1),
            "🔴 Impróprio (< 0.1)":        round((NDVI<0.1).mean()*100,1),
        }
        return NDVI, aptidao, pct

    def mostrar_resultados(img_rgb, RED, NIR, titulo=""):
        NDVI, aptidao, pct = calcular_ndvi_e_diagnostico(RED, NIR)

        col1,col2,col3,col4 = st.columns(4)
        with col1: st.metric("NDVI Médio", f"{NDVI.mean():.3f}")
        with col2: st.metric("NDVI Máx",   f"{NDVI.max():.3f}")
        with col3: st.metric("NDVI Mín",   f"{NDVI.min():.3f}")
        with col4: st.metric("✅ % Apto",  f"{pct['✅ Apto (NDVI ≥ 0.5)']:.1f}%")

        col_a, col_b = st.columns(2)
        with col_a:
            fig, axes = plt.subplots(1, 2, figsize=(10, 4))
            fig.patch.set_facecolor("#0d1117")
            if img_rgb is not None:
                axes[0].imshow(img_rgb)
                axes[0].set_title("Original RGB", color="white", fontsize=9)
            else:
                axes[0].imshow(RED, cmap="Reds")
                axes[0].set_title("Banda RED", color="white", fontsize=9)
            axes[0].axis("off")
            im = axes[1].imshow(NDVI, cmap="RdYlGn", vmin=-0.2, vmax=0.8)
            axes[1].set_title("Mapa NDVI", color="white", fontsize=9)
            axes[1].axis("off")
            plt.colorbar(im, ax=axes[1], fraction=0.046).ax.tick_params(colors="white",labelsize=7)
            st.pyplot(fig, use_container_width=True)

        with col_b:
            df_pct = pd.DataFrame({"Categoria":list(pct.keys()),"Percentual (%)":list(pct.values())})
            fig2 = px.pie(df_pct, values="Percentual (%)", names="Categoria",
                          color_discrete_sequence=["#27ae60","#f1c40f","#e67e22","#c0392b"],
                          template="plotly_dark", title="Distribuição de Aptidão Agrícola", hole=0.35)
            fig2.update_layout(margin=dict(l=0,r=0,t=40,b=0), height=300)
            st.plotly_chart(fig2, use_container_width=True)

        st.markdown("#### 📋 Diagnóstico Agronômico")
        apto = pct["✅ Apto (NDVI ≥ 0.5)"]
        mod  = pct["🟡 Moderado (0.3–0.5)"]
        if apto >= 60:
            st.success(f"✅ **SOLO PRONTO PARA AGRICULTURA** — {apto:.1f}% da área com vegetação saudável. Condições ideais de plantio.")
        elif apto + mod >= 50:
            st.warning(f"⚠️ **SOLO PARCIALMENTE APTO** — {apto+mod:.1f}% com cobertura razoável. Recomenda-se adubação e monitoramento.")
        else:
            st.error(f"❌ **SOLO IMPRÓPRIO / NECESSITA INTERVENÇÃO** — apenas {apto:.1f}% apto. Verificar irrigação, pragas e fertilidade.")

        with st.expander("📘 Ver recomendações detalhadas"):
            if apto >= 60:
                st.markdown("- Manter manejo atual\n- Monitorar NDVI mensalmente via satélite\n- Planejar colheita com base na produtividade estimada")
            elif apto + mod >= 50:
                st.markdown("- Aplicar adubação de cobertura nas áreas moderadas\n- Ajustar irrigação — verificar balanço hídrico\n- Nova análise em 15 dias")
            else:
                st.markdown("- Coletar amostras de solo para análise laboratorial\n- Aplicar defensivo localizado com drone\n- Verificar sistema de irrigação\n- Considerar plantio de cobertura")

    if modo == "📂 Upload de imagem real":
        uploaded = st.file_uploader("📁 Selecione uma ou mais imagens (JPG, PNG)",
                                    type=["jpg","jpeg","png"], accept_multiple_files=True)
        if uploaded:
            for arq in uploaded:
                st.markdown(f"#### 🖼️ {arq.name}")
                img = Image.open(arq).convert("RGB")
                img_array = np.array(img) / 255.0
                RED = img_array[:,:,0]
                NIR = img_array[:,:,1]
                mostrar_resultados(img_array, RED, NIR, arq.name)
                st.markdown("---")
        else:
            st.info("👆 Faça upload de uma imagem agrícola para iniciar a análise.")

    else:
        st.info("🔬 Usando imagem sintética com 4 zonas agrícolas distintas.")
        np.random.seed(42)
        SIZE = 100
        RED = np.zeros((SIZE,SIZE)); NIR = np.zeros((SIZE,SIZE))
        RED[:50,:50] = np.random.normal(0.08,0.015,(50,50)).clip(0,1)
        NIR[:50,:50] = np.random.normal(0.72,0.06,(50,50)).clip(0,1)
        RED[:50,50:] = np.random.normal(0.18,0.02,(50,50)).clip(0,1)
        NIR[:50,50:] = np.random.normal(0.50,0.07,(50,50)).clip(0,1)
        RED[50:,:50] = np.random.normal(0.30,0.03,(50,50)).clip(0,1)
        NIR[50:,:50] = np.random.normal(0.32,0.05,(50,50)).clip(0,1)
        RED[50:,50:] = np.random.normal(0.40,0.04,(50,50)).clip(0,1)
        NIR[50:,50:] = np.random.normal(0.22,0.03,(50,50)).clip(0,1)
        mostrar_resultados(None, RED, NIR, "Simulação")
